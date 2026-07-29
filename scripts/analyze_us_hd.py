"""Analyze US HD files to find intro→game transition points."""
import docx, re, os

base = r'C:\Users\鸢尾\Desktop\腾讯-模拟用户\资料\虚拟用户-笔录 for 元培\虚拟用户-笔录 for 元培\美国HD端射击市场用户细分研究'

for fname in sorted(os.listdir(base)):
    if not fname.endswith('.docx'):
        continue
    doc = docx.Document(os.path.join(base, fname))

    # Find moderator
    mod_num = None
    for i, p in enumerate(doc.paragraphs[:100]):
        text = p.text.strip()
        if re.search(r'(?i)(my name is moderator|independent moderator|i.?\s*m an? independent|moderator.*speaking)', text):
            for j in range(i - 1, max(i - 3, -1), -1):
                prev = doc.paragraphs[j].text.strip()
                m = re.match(r'SPEAKER_(\d+)\(', prev)
                if m:
                    mod_num = m.group(1)
                    break
            if mod_num:
                break

    # Parse turns
    turns = []
    cur_spk = None; cur_txt = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        m = re.match(r'SPEAKER_(\d+)\([^)]*\):\s*$', text)
        if m:
            if cur_spk and cur_txt and cur_txt.strip():
                turns.append((cur_spk, cur_txt.strip()))
            cur_spk = m.group(1); cur_txt = None; continue
        m = re.match(r'SPEAKER_(\d+)\([^)]*\):\s*(.+)', text)
        if m:
            if cur_spk and cur_txt and cur_txt.strip():
                turns.append((cur_spk, cur_txt.strip()))
            cur_spk = m.group(1); cur_txt = m.group(2); continue
        if re.match(r'\([^)]*\):\s*', text):
            continue
        if cur_spk:
            cur_txt = (cur_txt + ' ' + text) if cur_txt else text
    if cur_spk and cur_txt and cur_txt.strip():
        turns.append((cur_spk, cur_txt.strip()))

    # Show moderator turns to find transition
    mod_turns = [(i, t) for i, (s, t) in enumerate(turns) if s == mod_num]
    print(f'\n=== {fname} (mod=SPEAKER_{mod_num}, {len(turns)} turns) ===')
    for i, t in mod_turns[:25]:
        print(f'  Turn {i}: {t[:200]}')
    if len(mod_turns) > 25:
        print(f'  ... ({len(mod_turns)} mod turns total)')