#!/usr/bin/env python3
"""Extract segments from 漫威争锋中美用户洞察研究 docx files to JSON."""

import docx, os, json, re

base = r'C:\Users\鸢尾\Desktop\腾讯-模拟用户\资料\虚拟用户-笔录 for 元培\虚拟用户-笔录 for 元培\漫威争锋中美用户洞察研究'
outdir = r'data\群体画像'

# ============================================================
# FILE 1: 座谈会笔录-G1.docx (table format)
# ============================================================
print('=== Processing 座谈会笔录-G1.docx ===')
doc = docx.Document(os.path.join(base, '座谈会笔录-G1.docx'))

# Step 1: Parse speaker mapping from header paragraphs
# [3] 1G1-LJH, [4] 2张嘉尹, [5] 3G1-LWX, [6] 4王浩东, [7] 5G1-MXC, [8] 6郑耿麟, [9] 7G1-ZHS, [10] 8陈志安
speaker_map = {}
for p in doc.paragraphs:
    text = p.text.strip()
    # Match: "1G1-LJH" or "2张嘉尹" etc.
    m = re.match(r'(\d+)\s*(.+)', text)
    if m:
        num = m.group(1)
        name = m.group(2).strip()
        # Accept names that are reasonable length (not full sentences)
        if name and len(name) < 20 and not name.startswith('实录'):
            speaker_map[num] = name
            print(f'  Header: {num} -> {name}')

# If header parsing didn't work well, try to extract from first table rows
if len(speaker_map) < 5:
    print('  Header parsing insufficient, trying table-based extraction...')
    table = doc.tables[0]
    for row in table.rows[:30]:
        col1 = row.cells[1].text.strip()
        # Match: '1：大家好，我叫G1-LJH...' or '2：大家好，我叫张嘉尹...'
        m = re.match(r'(\d+)[：:]\s*大家好[，,]?\s*[我我叫是]+\s*(\S{1,10})', col1)
        if m:
            num = m.group(1)
            name = m.group(2).strip('，,。. ')
            if num not in speaker_map:
                speaker_map[num] = name
                print(f'  Table: {num} -> {name}')

# Fix overly long names by extracting G1-xxx IDs from the table
table = doc.tables[0]
for row in table.rows[:30]:
    col1 = row.cells[1].text.strip()
    m = re.search(r'(G1-\w+)', col1)
    if m:
        gid = m.group(1)
        num_m = re.match(r'(\d+)[：:]', col1)
        if num_m:
            num = num_m.group(1)
            if num in speaker_map:
                old = speaker_map[num]
                if len(old) > 10:
                    speaker_map[num] = gid
                    print(f'  Fixed: {num} -> {gid} (was too long)')

print(f'  Final speaker_map: {speaker_map}')

# Step 2: Extract segments from table
segments_g1 = []
last_mod_text = None
seg_idx = 0

for row in table.rows:
    mod_text = row.cells[0].text.strip()
    int_text = row.cells[1].text.strip()

    # Process moderator text
    if mod_text:
        # Clean moderator prefix "M：" or "M:"
        mod_clean = re.sub(r'^[Mm]\s*[：:]\s*', '', mod_text).strip()
        # Only update if it's a real question (not just a name callout)
        # Filter out: pure name, name+单字, name+说一下/介绍一下 etc.
        is_name_only = (
            len(mod_clean) <= 10 or
            re.match(r'^[一-鿿A-Za-z0-9_-]{1,15}[。，,.]?\s*$', mod_clean) or
            re.match(r'^[一-鿿A-Za-z0-9_-]{1,10}[说一下介绍一下][。，,.]?\s*$', mod_clean)
        )
        if mod_clean and not is_name_only:
            last_mod_text = mod_clean

    if not int_text:
        continue

    # Parse interviewee: format like '1：text...' or 'G1-LJH：text...'
    m = re.match(r'(\d+)\s*[：:]\s*(.+)', int_text, re.DOTALL)
    if m:
        num = m.group(1)
        text = m.group(2).strip()
        speaker_id = speaker_map.get(num, f'P{num}')
    else:
        # Try G1-xxx： format
        m = re.match(r'(G1-\w+)\s*[：:]\s*(.+)', int_text, re.DOTALL)
        if m:
            speaker_id = m.group(1)
            text = m.group(2).strip()
        else:
            continue

    if not text or len(text) < 2:
        continue

    segments_g1.append({
        'source_file': '漫威争锋中美用户洞察研究/座谈会笔录-G1.docx',
        'segment_index': seg_idx,
        'speaker_id': speaker_id,
        'speaker_role': 'interviewee',
        'preceding_question': last_mod_text,
        'original_text': text,
        'cleaned_text': None,
        'char_count': None,
        'annotation': {}
    })
    seg_idx += 1

outpath = os.path.join(outdir, 'segments_漫威争锋_座谈会笔录-G1.json')
with open(outpath, 'w', encoding='utf-8') as f:
    json.dump(segments_g1, f, ensure_ascii=False, indent=2)
print(f'  Saved {len(segments_g1)} segments')

# Verify first 10
for i in range(min(10, len(segments_g1))):
    s = segments_g1[i]
    pq = 'YES' if s['preceding_question'] else 'none'
    print(f'  [{i}] speaker={s["speaker_id"]}, pq={pq}, text={s["original_text"][:80]}...')

# ============================================================
# FILE 2 & 3: 文字转录 (paragraph format)
# ============================================================
for fname in ['文字转录-海外大龄组.docx', '文字转录-海外年轻组.docx']:
    print(f'\n=== Processing {fname} ===')
    doc = docx.Document(os.path.join(base, fname))

    segments = []
    last_mod_text = None
    seg_idx = 0

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # Parse format: STF#说话人XX(timestamp): text
        m = re.match(r'STF#说话人(\d+)\([^)]+\):\s*(.+)', text, re.DOTALL)
        if not m:
            # Might be continuation of previous speech - append to last segment
            if segments:
                segments[-1]['original_text'] += '\n' + text
            continue

        speaker_num = m.group(1)
        speech = m.group(2).strip()

        if speaker_num == '01':
            # Moderator
            last_mod_text = speech
        else:
            # Interviewee
            speaker_id = f'STF{speaker_num}'
            segments.append({
                'source_file': f'漫威争锋中美用户洞察研究/{fname}',
                'segment_index': seg_idx,
                'speaker_id': speaker_id,
                'speaker_role': 'interviewee',
                'preceding_question': last_mod_text,
                'original_text': speech,
                'cleaned_text': None,
                'char_count': None,
                'annotation': {}
            })
            seg_idx += 1

    base_name = fname.replace('.docx', '')
    outpath = os.path.join(outdir, f'segments_漫威争锋_{base_name}.json')
    with open(outpath, 'w', encoding='utf-8') as f:
        json.dump(segments, f, ensure_ascii=False, indent=2)
    print(f'  Saved {len(segments)} segments')

    # Verify first 5
    for i in range(min(5, len(segments))):
        s = segments[i]
        pq = 'YES' if s['preceding_question'] else 'none'
        print(f'  [{i}] speaker={s["speaker_id"]}, pq={pq}, text={s["original_text"][:80]}...')

print('\n=== All done! ===')