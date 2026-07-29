"""Extract ALL 漫威争锋 files: respondents + segments."""
import docx, re, os, json

base = r'C:\Users\鸢尾\Desktop\腾讯-模拟用户\资料\虚拟用户-笔录 for 元培\虚拟用户-笔录 for 元培\漫威争锋中美用户洞察研究'
outdir = r'data\sheets_data'

NOISE = re.compile(
    r'^[对是嗯好行可可以]+[，,。.]?$|'
    r'^[没不][有会是知道清楚懂行能]+[，,。.]?$|'
    r'^(Cool|Yeah|Yep|Nope|Yes|No|Okay|Ok|Sure|Right|Gotcha|Perfect|Fantastic|Great|Nice|Thanks|Sorry|Fine|Alright|Absolutely|Exactly|Definitely|Totally|Indeed|Correct|Fair|True|Maybe|Perhaps|Probably)$|'
    r'^[Mm]+[hm]+[，,。.]?$|'
    r'^[啊哦嗯呃唉哎哟嘿]$',
    re.IGNORECASE
)

# ============================================================
# FILE 1: 座谈会笔录-G1.docx (table format, Chinese)
# ============================================================
print('=== 座谈会笔录-G1.docx ===')
doc = docx.Document(os.path.join(base, '座谈会笔录-G1.docx'))

# Speaker map from header
speaker_map = {}
for p in doc.paragraphs:
    text = p.text.strip()
    m = re.match(r'(\d+)\s*(\S.+)', text)
    if m and len(m.group(2)) < 20 and not m.group(2).startswith('实录'):
        speaker_map[m.group(1)] = m.group(2).strip()

# Parse all table rows
all_rows = []
table = doc.tables[0]
for row in table.rows:
    mod = row.cells[0].text.strip()
    intv = row.cells[1].text.strip()
    if intv:
        m = re.match(r'(\d+)\s*[：:]\s*(.+)', intv, re.DOTALL)
        if m:
            num, text = m.group(1), m.group(2).strip()
            sid = speaker_map.get(num, f'P{num}')
            mod_clean = re.sub(r'^[Mm]\s*[：:]\s*', '', mod).strip() if mod else ''
            all_rows.append({'speaker_id': sid, 'mod': mod_clean, 'text': text, 'num': num})

# Respondents: rows 0-8 (self-introductions)
respondents = {}
for r in all_rows[:9]:
    sid = r['speaker_id']
    if sid not in respondents:
        respondents[sid] = {
            'source_file': '漫威争锋中美用户洞察研究/座谈会笔录-G1.docx',
            'speaker_id': sid, 'display_name': None, 'group_code': 'G1',
            'background': {'profile': '', 'game_experience': None, 'game_experience_summary': None}
        }
    respondents[sid]['background']['profile'] += ('；' if respondents[sid]['background']['profile'] else '') + r['text']

# Segments: rows 9+ (game discussion)
segments = []
last_mod = None
seg_idx = 0
for r in all_rows[9:]:
    if r['mod'] and len(r['mod']) > 10:
        if not re.match(r'^[一-鿿A-Za-z0-9_\-]{1,10}[说一下介绍一下补充一下来说说。，,.]?\s*$', r['mod']):
            last_mod = r['mod']
    if NOISE.match(r['text']) or len(r['text']) <= 2:
        continue
    segments.append({
        'source_file': '漫威争锋中美用户洞察研究/座谈会笔录-G1.docx',
        'segment_index': seg_idx, 'speaker_id': r['speaker_id'], 'speaker_role': 'interviewee',
        'preceding_question': last_mod, 'original_text': r['text'],
        'cleaned_text': None, 'char_count': None, 'annotation': {}
    })
    seg_idx += 1

# Save
prefix = '漫威争锋中美用户洞察研究'
with open(os.path.join(outdir, f'respondents_{prefix}_座谈会笔录-G1.json'), 'w', encoding='utf-8') as f:
    json.dump(list(respondents.values()), f, ensure_ascii=False, indent=2)
with open(os.path.join(outdir, f'segments_{prefix}_座谈会笔录-G1.json'), 'w', encoding='utf-8') as f:
    json.dump(segments, f, ensure_ascii=False, indent=2)
print(f'  respondents: {len(respondents)}, segments: {len(segments)}')

# ============================================================
# FILE 2 & 3: 文字转录 (paragraph format, English)
# ============================================================
# Manual intro_end based on content analysis:
# - 海外年轻组: moderator (STF01=Riley S.) asks intros turns 7-38, game starts turn 40
# - 海外大龄组: casual chat format, no structured intro, all game discussion
FILE_INTRO_END = {
    '文字转录-海外年轻组.docx': 40,
    '文字转录-海外大龄组.docx': 0,  # No intro section
}

for fname in ['文字转录-海外大龄组.docx', '文字转录-海外年轻组.docx']:
    print(f'=== {fname} ===')
    doc = docx.Document(os.path.join(base, fname))

    # Parse all turns
    turns = []
    current_speaker = None
    current_text = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # New speaker line (SPEAKER only, no text)
        m = re.match(r'STF#说话人(\d+)\([^)]*\):\s*$', text)
        if m:
            if current_speaker is not None and current_text and current_text.strip():
                turns.append({'speaker': current_speaker, 'text': current_text.strip()})
            current_speaker = m.group(1)
            current_text = None
            continue

        # New speaker line with inline text
        m = re.match(r'STF#说话人(\d+)\([^)]*\):\s*(.+)', text)
        if m:
            if current_speaker is not None and current_text and current_text.strip():
                turns.append({'speaker': current_speaker, 'text': current_text.strip()})
            current_speaker = m.group(1)
            current_text = m.group(2)
            continue

        # Skip anonymous timestamp lines
        if re.match(r'\([^)]*\):\s*', text):
            continue

        # Continuation
        if current_speaker is not None:
            current_text = (current_text + ' ' + text) if current_text else text

    # Save last turn
    if current_speaker is not None and current_text and current_text.strip():
        turns.append({'speaker': current_speaker, 'text': current_text.strip()})

    # Use manual intro_end for this file
    intro_end = FILE_INTRO_END.get(fname, 0)

    # Extract respondents from intro section
    respondents = {}
    for t in turns[:intro_end]:
        if t['speaker'] == '01':
            continue  # Skip moderator
        sid = f'STF{t["speaker"]}'
        if sid not in respondents:
            respondents[sid] = {
                'source_file': f'漫威争锋中美用户洞察研究/{fname}',
                'speaker_id': sid, 'display_name': None, 'group_code': None,
                'background': {'profile': '', 'game_experience': None, 'game_experience_summary': None}
            }
        respondents[sid]['background']['profile'] += ('；' if respondents[sid]['background']['profile'] else '') + t['text']

    # Extract segments from game discussion
    segments = []
    last_mod = None
    seg_idx = 0
    for t in turns[intro_end:]:
        if t['speaker'] == '01':
            last_mod = t['text']
        else:
            if NOISE.match(t['text']) or len(t['text']) <= 2:
                continue
            segments.append({
                'source_file': f'漫威争锋中美用户洞察研究/{fname}',
                'segment_index': seg_idx,
                'speaker_id': f'STF{t["speaker"]}',
                'speaker_role': 'interviewee',
                'preceding_question': last_mod,
                'original_text': t['text'],
                'cleaned_text': None, 'char_count': None, 'annotation': {}
            })
            seg_idx += 1

    # Save
    base_name = fname.replace('.docx', '')
    with open(os.path.join(outdir, f'respondents_{prefix}_{base_name}.json'), 'w', encoding='utf-8') as f:
        json.dump(list(respondents.values()), f, ensure_ascii=False, indent=2)
    with open(os.path.join(outdir, f'segments_{prefix}_{base_name}.json'), 'w', encoding='utf-8') as f:
        json.dump(segments, f, ensure_ascii=False, indent=2)
    print(f'  respondents: {len(respondents)}, segments: {len(segments)} (intro_end={intro_end})')

print('\n=== All done! ===')