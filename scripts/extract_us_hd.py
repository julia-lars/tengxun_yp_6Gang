#!/usr/bin/env python3
"""Extract segments from 美国HD端射击市场用户细分研究 docx files to JSON."""

import docx, os, json, re

base = r'C:\Users\鸢尾\Desktop\腾讯-模拟用户\资料\虚拟用户-笔录 for 元培\虚拟用户-笔录 for 元培\美国HD端射击市场用户细分研究'
outdir = r'data\sheets_data'

for fname in sorted(os.listdir(base)):
    if not fname.endswith('.docx'):
        continue

    print(f'=== Processing {fname} ===')
    full = os.path.join(base, fname)
    doc = docx.Document(full)

    # Step 1: Detect moderator speaker ID
    mod_speaker_num = None
    for i, p in enumerate(doc.paragraphs[:100]):
        text = p.text.strip()
        if re.search(r'(?i)(my name is moderator|independent moderator|i\Wm an? independent|moderator[ ,].*speaking|I\Wm.*moderator)', text):
            # Look back for the speaker ID
            for j in range(i - 1, max(i - 3, -1), -1):
                prev = doc.paragraphs[j].text.strip()
                m = re.match(r'SPEAKER_(\d+)\(', prev)
                if m:
                    mod_speaker_num = m.group(1)
                    break
            if mod_speaker_num:
                break

    # Fallback: try to find moderator by checking for "Moderator" keyword in any text
    if not mod_speaker_num:
        for i, p in enumerate(doc.paragraphs[:200]):
            text = p.text.strip()
            if re.search(r'(?i)\bmoderator\b', text):
                for j in range(i - 1, max(i - 3, -1), -1):
                    prev = doc.paragraphs[j].text.strip()
                    m = re.match(r'SPEAKER_(\d+)\(', prev)
                    if m:
                        mod_speaker_num = m.group(1)
                        break
                if mod_speaker_num:
                    break

    if not mod_speaker_num:
        # Last resort: assume the most frequent speaker in first 100 paragraphs is moderator
        from collections import Counter
        speaker_counts = Counter()
        for p in doc.paragraphs[:100]:
            m = re.match(r'SPEAKER_(\d+)\(', p.text.strip())
            if m:
                speaker_counts[m.group(1)] += 1
        if speaker_counts:
            mod_speaker_num = speaker_counts.most_common(1)[0][0]

    print(f'  Moderator: SPEAKER_{mod_speaker_num}')

    # Step 2: Parse paragraphs into speaker turns
    current_speaker = None
    current_text = None
    last_mod_text = None
    segments = []
    seg_idx = 0

    for p in doc.paragraphs:
        text = p.text.strip()

        # Check if this is a new speaker line
        m = re.match(r'SPEAKER_(\d+)\([^)]*\):\s*$', text)
        if m:
            # Save previous turn
            if current_speaker is not None and current_text and current_text.strip():
                if current_speaker == mod_speaker_num:
                    last_mod_text = current_text.strip()
                else:
                    speaker_id = f'SPEAKER_{current_speaker}'
                    segments.append({
                        'source_file': f'美国HD端射击市场用户细分研究/{fname}',
                        'segment_index': seg_idx,
                        'speaker_id': speaker_id,
                        'speaker_role': 'interviewee',
                        'preceding_question': last_mod_text,
                        'original_text': current_text.strip(),
                        'cleaned_text': None,
                        'char_count': None,
                        'annotation': {}
                    })
                    seg_idx += 1

            # Start new turn
            current_speaker = m.group(1)
            current_text = None
            continue

        # Also match: SPEAKER_XX(timestamp): text on same line
        m = re.match(r'SPEAKER_(\d+)\([^)]*\):\s*(.+)', text)
        if m:
            # Save previous turn
            if current_speaker is not None and current_text and current_text.strip():
                if current_speaker == mod_speaker_num:
                    last_mod_text = current_text.strip()
                else:
                    segments.append({
                        'source_file': f'美国HD端射击市场用户细分研究/{fname}',
                        'segment_index': seg_idx,
                        'speaker_id': f'SPEAKER_{current_speaker}',
                        'speaker_role': 'interviewee',
                        'preceding_question': last_mod_text,
                        'original_text': current_text.strip(),
                        'cleaned_text': None,
                        'char_count': None,
                        'annotation': {}
                    })
                    seg_idx += 1

            # Start new turn with inline text
            current_speaker = m.group(1)
            current_text = m.group(2)
            continue

        # Skip empty lines
        if not text:
            continue

        # Skip lines without speaker prefix (staff/admin lines like "(00:00:16):")
        if re.match(r'\([^)]*\):\s*$', text) or re.match(r'\([^)]*\):\s*.+', text):
            continue

        # Continuation of current speaker's text
        if current_speaker is not None:
            if current_text is None:
                current_text = text
            else:
                current_text += ' ' + text

    # Save last turn
    if current_speaker is not None and current_text and current_text.strip():
        if current_speaker == mod_speaker_num:
            pass  # Don't save moderator's last turn
        else:
            segments.append({
                'source_file': f'美国HD端射击市场用户细分研究/{fname}',
                'segment_index': seg_idx,
                'speaker_id': f'SPEAKER_{current_speaker}',
                'speaker_role': 'interviewee',
                'preceding_question': last_mod_text,
                'original_text': current_text.strip(),
                'cleaned_text': None,
                'char_count': None,
                'annotation': {}
            })
            seg_idx += 1

    # Save
    base_name = fname.replace('.docx', '')
    outpath = os.path.join(outdir, f'segments_美国HD端_{base_name}.json')
    with open(outpath, 'w', encoding='utf-8') as f:
        json.dump(segments, f, ensure_ascii=False, indent=2)
    print(f'  Saved {len(segments)} segments')

    # Verify first 3
    for i in range(min(3, len(segments))):
        s = segments[i]
        pq = 'YES' if s['preceding_question'] else 'none'
        txt = s['original_text'][:100]
        print(f'  [{i}] speaker={s["speaker_id"]}, pq={pq}, text={txt}...')

print('\n=== All done! ===')