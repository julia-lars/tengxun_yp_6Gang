"""Extract US HD端射击市场用户细分研究: respondents + segments, merge into 2 files."""
import docx, re, os, json

base = r'C:\Users\鸢尾\Desktop\腾讯-模拟用户\资料\虚拟用户-笔录 for 元培\虚拟用户-笔录 for 元培\美国HD端射击市场用户细分研究'
outdir = r'data\群体画像'
prefix = '美国HD端射击市场用户细分研究'

# Manual intro_end based on moderator transition analysis
FILE_INTRO_END = {
    '座谈会笔录-G1.docx': 36,
    '座谈会笔录-G2.docx': 26,
    '座谈会笔录-G3.docx': 39,
    '座谈会笔录-G4.docx': 36,
    '座谈会笔录-G5.docx': 27,
    '座谈会笔录-G6.docx': 15,
    '座谈会笔录-G7.docx': 23,
    '座谈会笔录-G8.docx': 15,
}

NOISE = re.compile(
    r'^[对是嗯好行可可以]+[，,。.]?$|'
    r'^[没不][有会是知道清楚懂行能]+[，,。.]?$|'
    r'^(Cool|Yeah|Yep|Nope|Yes|No|Okay|Ok|Sure|Right|Gotcha|Perfect|Fantastic|Great|Nice|Thanks|Sorry|Fine|Alright|Absolutely|Exactly|Definitely|Totally|Indeed|Correct|Fair|True|Maybe|Perhaps|Probably)$|'
    r'^[Mm]+[hm]+[，,。.]?$|'
    r'^[啊哦嗯呃唉哎哟嘿]$',
    re.IGNORECASE
)

all_respondents = []
all_segments = []

for fname in sorted(os.listdir(base)):
    if not fname.endswith('.docx'):
        continue
    print(f'=== {fname} ===')
    doc = docx.Document(os.path.join(base, fname))

    # Find moderator
    mod_num = None
    for i, p in enumerate(doc.paragraphs[:100]):
        text = p.text.strip()
        if re.search(r'(?i)(my name is moderator|independent moderator|i.?\s*m an? independent|moderator.*speaking)', text):
            for j in range(i - 1, max(i - 3, -1), -1):
                prev = doc.paragraphs[j].text.strip()
                m = re.match(r'SPEAKER_(\d+)\(', prev)
                if m:
                    mod_num = m.group(1)
                    break
            if mod_num:
                break
    if not mod_num:
        print(f'  WARNING: moderator not found, skipping')
        continue

    # Parse turns
    turns = []
    cur_spk = None; cur_txt = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        m = re.match(r'SPEAKER_(\d+)\([^)]*\):\s*$', text)
        if m:
            if cur_spk and cur_txt and cur_txt.strip():
                turns.append({'speaker': cur_spk, 'text': cur_txt.strip()})
            cur_spk = m.group(1); cur_txt = None; continue
        m = re.match(r'SPEAKER_(\d+)\([^)]*\):\s*(.+)', text)
        if m:
            if cur_spk and cur_txt and cur_txt.strip():
                turns.append({'speaker': cur_spk, 'text': cur_txt.strip()})
            cur_spk = m.group(1); cur_txt = m.group(2); continue
        if re.match(r'\([^)]*\):\s*', text):
            continue
        if cur_spk:
            cur_txt = (cur_txt + ' ' + text) if cur_txt else text
    if cur_spk and cur_txt and cur_txt.strip():
        turns.append({'speaker': cur_spk, 'text': cur_txt.strip()})

    intro_end = FILE_INTRO_END.get(fname, 0)

    # Extract respondents from intro
    respondents = {}
    for t in turns[:intro_end]:
        if t['speaker'] == mod_num:
            continue
        sid = f'SPEAKER_{t["speaker"]}'
        if sid not in respondents:
            respondents[sid] = {
                'source_file': f'{prefix}/{fname}',
                'speaker_id': sid, 'display_name': None, 'group_code': None,
                'background': {'profile': '', 'game_experience': None, 'game_experience_summary': None}
            }
        respondents[sid]['background']['profile'] += ('; ' if respondents[sid]['background']['profile'] else '') + t['text']

    # Extract segments from game discussion
    segments = []
    last_mod = None
    seg_idx = 0
    for t in turns[intro_end:]:
        if t['speaker'] == mod_num:
            last_mod = t['text']
        else:
            if NOISE.match(t['text']) or len(t['text']) <= 2:
                continue
            segments.append({
                'source_file': f'{prefix}/{fname}',
                'segment_index': seg_idx,
                'speaker_id': f'SPEAKER_{t["speaker"]}',
                'speaker_role': 'interviewee',
                'preceding_question': last_mod,
                'original_text': t['text'],
                'cleaned_text': None, 'char_count': None, 'annotation': {}
            })
            seg_idx += 1

    all_respondents.extend(respondents.values())
    all_segments.extend(segments)
    print(f'  respondents: {len(respondents)}, segments: {len(segments)}')

# Save merged files
out_r = os.path.join(outdir, f'respondents_{prefix}.json')
with open(out_r, 'w', encoding='utf-8') as f:
    json.dump(all_respondents, f, ensure_ascii=False, indent=2)
print(f'\nrespondents: {len(all_respondents)} total -> {out_r}')

out_s = os.path.join(outdir, f'segments_{prefix}.json')
with open(out_s, 'w', encoding='utf-8') as f:
    json.dump(all_segments, f, ensure_ascii=False, indent=2)
print(f'segments: {len(all_segments)} total -> {out_s}')
print('Done!')