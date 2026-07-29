#!/usr/bin/env python3
"""Preview G1 table rows to understand filtering needs."""
import docx, re, os

base = r'C:\Users\鸢尾\Desktop\腾讯-模拟用户\资料\虚拟用户-笔录 for 元培\虚拟用户-笔录 for 元培\漫威争锋中美用户洞察研究'

doc = docx.Document(os.path.join(base, '座谈会笔录-G1.docx'))
table = doc.tables[0]

# Header: build speaker map
speaker_map = {}
for p in doc.paragraphs:
    text = p.text.strip()
    m = re.match(r'(\d+)\s*(.+)', text)
    if m and len(m.group(2).strip()) < 20 and not m.group(2).startswith('实录'):
        speaker_map[m.group(1)] = m.group(2).strip()
print('Speaker map:', speaker_map)

# Print rows 0-80
for ri, row in enumerate(table.rows[:80]):
    mod = row.cells[0].text.strip()
    intv = row.cells[1].text.strip()
    if intv:
        m = re.match(r'(\d+)\s*[：:]\s*(.+)', intv, re.DOTALL)
        if m:
            num = m.group(1)
            text = m.group(2).strip()
            sid = speaker_map.get(num, f'P{num}')
            mod_preview = mod[:60] if mod else '(none)'
            print(f'Row {ri}: speaker={sid}, mod={mod_preview}, text={text[:150]}')
        else:
            print(f'Row {ri}: NO MATCH, intv={intv[:100]}')