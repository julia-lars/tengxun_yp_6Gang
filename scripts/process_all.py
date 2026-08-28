#!/usr/bin/env python3
"""
批量处理所有访谈文件 → 群体画像v2.2 (PQ 质量修复版)
修复 v2.1 的 PQ 质量问题：
  1. 主持人自动检测（替代硬编码 SPEAKER_05）
  2. 空 PQ 回填仅使用真实提问
  3. PQ 长度上限 200 字符
  4. 英文对话分类优化
  5. 文字转录特殊处理
保留 v2.0→v2.1 的修复：
  6. Gaming Background 从筛选表提取
  7. Profile/GB 从 Segment 文本提取（无筛选表文件）
  8. group_code 从文件路径提取
  9. 超短 Segment 合并（≤5字符）
  10. xlsx 按 §6.4 规范处理
  11. display_name 修复
"""

from docx import Document
from openpyxl import load_workbook
import json, re, os, sys
from collections import Counter, OrderedDict

BASE = "/Users/jessicajyan/Desktop/腾讯用户画像-data/data/虚拟用户-笔录 for 元培"
OUT = "/Users/jessicajyan/tengxun_yp_6Gang/data/群体画像v2.0"

MAX_PQ_LENGTH = 200  # §6.7 PQ 长度上限

# English patterns for classify_mod
EN_BANTER_RE = re.compile(
    r'^(yeah[,\s]|yep[,\s]|nope[,\s]|ok[,\s]|okay[,\s]|cool[,\s]|dude[,\s]|nice[,\s]|'
    r'great[,\s]|awesome[,\s]|fantastic[,\s]|perfect[,\s]|wow[,\s]|right[,\s]|'
    r'sure[,\s]|absolutely[,\s]|exactly[,\s]|totally[,\s]|fair[,\s]|true[,\s]|'
    r'gotcha[,\s]|alright[,\s]|sorry[,\s]|thanks[,\s]|thank you[,\s]|'
    r'what\'s up|whats up|howdy|hey there|'
    r'i think\b|i mean\b|i feel\b|i guess\b|i know\b|i\'ve\b|i\'m\b|i was\b|i had\b|'
    r'you know\b|it\'s like\b|that\'s\b|this is\b|'
    r'and then\b|but yeah\b|so yeah\b|so i\b|so we\b|also like\b)',
    re.IGNORECASE
)

EN_QUESTION_RE = re.compile(
    r'(how about you|what about you|what do you think|how do you feel|'
    r'how do you|tell me about|can you |could you |would you |'
    r'do you |have you |are you |is there |'
    r'what\'s your|what is your|what are your|what was your|'
    r'how would|how does|how is|how has|'
    r'any thoughts|any questions|what else|anything else|'
    r'why do you|why did you|why are you|'
    r'which one|which of|which game|'
    r'when did you|when was the|'
    r'who do you|who is)',
    re.IGNORECASE
)

EN_INSTRUCTION_RE = re.compile(
    r'(let\'s start|let\'s begin|let\'s get to know|let\'s go|'
    r'we\'re going to|today we\'ll|first of all|first thing|'
    r'introduce yourself|tell us about yourself|'
    r'go ahead and|feel free to|'
    r'we\'ll start|we\'re here|thanks for coming|'
    r'welcome|all right|we\'re recording|'
    r'let\'s do|let\'s talk|let\'s hear|'
    r'we\'re gonna|we\'re going to be)',
    re.IGNORECASE
)

# Moderator-like phrases for auto-detection
MOD_PHRASES = [
    "welcome", "thanks for coming", "today we're going", "all right",
    "let's get started", "we're recording", "my name is moderator",
    "we're here for", "we're going to be here", "let's start",
    "thank you for joining", "thanks for joining", "we're gonna",
    "let's go ahead", "we'll start", "we'll be here",
]

# ============================================================
# Core helpers
# ============================================================

def extract_core_question(text):
    text = text.strip()
    if len(text) <= 60:
        if len(text) > MAX_PQ_LENGTH:
            return text[:MAX_PQ_LENGTH] + "…"
        return text
    parts = re.split(r'([。！？；])', text)
    sentences = []
    for i in range(0, len(parts) - 1, 2):
        sent = parts[i] + parts[i+1]
        if sent.strip():
            sentences.append(sent.strip())
    if len(parts) % 2 == 1 and parts[-1].strip():
        sentences.append(parts[-1].strip())
    if len(sentences) > 1:
        candidate = sentences[-1]
        if len(candidate) < 20 and len(sentences) >= 2:
            candidate = sentences[-2] + candidate
    else:
        candidate = text
    if len(candidate) > 80:
        clauses = re.split(r'[，,]+', candidate)
        clauses = [c.strip() for c in clauses if c.strip() and len(c.strip()) >= 4]
        if len(clauses) > 4:
            candidate = '，'.join(clauses[-4:])
    if len(candidate) > MAX_PQ_LENGTH:
        candidate = candidate[:MAX_PQ_LENGTH] + "…"
    return candidate


def classify_mod(text):
    text = text.strip()
    if not text:
        return ("confirm", text)
    core = extract_core_question(text)

    name_call_re = re.compile(r'^好?，?\s*G\d+-?[A-Za-z0-9_]*\s*[呢吧]?[。？]?\s*$')
    if name_call_re.match(text):
        return ("name_call", core)
    if text in ("下一个。", "好，下一个。", "下一个。", "Next.", "next."):
        return ("name_call", text)

    name_call_end_re = re.compile(r'G\d+-?[A-Za-z0-9_]*\s*[呢吧]?[。？]?\s*$')
    if name_call_end_re.search(text):
        if not re.match(r'^G\d+-?[A-Za-z0-9_]*\s*[，,]\s*', text) and not re.match(r'^G\d+-?[A-Za-z0-9_]*\s*你', text):
            return ("name_call", core)

    # English: detect banter/chitchat → confirm
    if EN_BANTER_RE.match(text):
        return ("confirm", text)

    # English: detect instructions / topic transitions → new_topic
    if EN_INSTRUCTION_RE.search(text):
        return ("new_topic", core)

    # English: detect real questions → direct_question
    if EN_QUESTION_RE.search(text):
        return ("direct_question", core)

    if len(text) < 30 and '？' not in text and '?' not in text:
        confirm_patterns = (
            r'^(对|嗯|好|没问题|勉强也算|综合的|杂食|没有|也可以|差不多|算是吧|就是|'
            r'你也没有|简单来讲|后面|也是|好像|其实|所以|没问题|勉强|加一|差不多|'
            r'不管|那|这种|这些|那个|可以|OK|ok|确实|之前|玩过|之前玩|之前玩的|'
            r'之前玩过|之前玩过一段|之前玩过一段时间|你之前玩过|我看你|成长|回到|因为|'
            r'你觉得|你也觉得|但是|还有|可能|我觉|听起来|我理解|它也|'
            r'所以整体|看起来|没问题|好|OK|Alright|Cool|Yeah|Right|Okay|Sure|'
            r'Right|Great|Good|Interesting|Awesome|Wonderful|Got it|I see|'
            r'And|Yeah|Okay|Alright|Um|Hmm|So|Also|But|Like|'
            r'That makes sense|Fair enough|Good to know|'
            r'And then|So then|What about|How about|What else|Anything else|'
            r'Tell me more|Go on|Continue|'
            r'Any other|Anything|Anyone|Somebody|Somebody else)'
        )
        if re.match(confirm_patterns, text, re.IGNORECASE):
            return ("confirm", text)

    new_topic_markers = [
        "我们进入今天的正体", "我们进入下一部分", "下一部分进入到",
        "我们聊聊", "大家先填一下", "大家快速的说一下",
        "Let's get started", "Let's begin", "First of all",
        "I'd like to start", "To begin with", "We'll start",
    ]
    if any(m in text for m in new_topic_markers):
        return ("new_topic", core)

    name_q_re = re.compile(r'^G\d+-?[A-Za-z0-9_]*\s*[，,]\s*')
    name_q_ni_re = re.compile(r'^G\d+-?[A-Za-z0-9_]*\s*你')
    if name_q_re.match(text) or name_q_ni_re.match(text):
        return ("direct_question", core)

    group_words = ["大家", "每个人", "你们", "各位", "everyone", "everybody", "anybody"]
    directive_words = ["介绍一下", "讲一下", "聊一下", "说一下", "谈一下", "分享一下",
                       "introduce", "describe", "share", "tell me", "explain"]
    if len(text) > 50:
        if any(p in text for p in group_words) and any(p in text for p in directive_words):
            return ("new_topic", core)

    if len(core) <= 15 and ('？' in text or '?' in text):
        return ("true_followup", core)

    directive_phrases = ["来反驳", "来说说", "来讲讲", "来补充", "来回应", "说说看", "继续讲", "讲一下", "说一下"]
    if any(p in text for p in directive_phrases):
        return ("direct_question", core)

    if len(text) > 80 and ('？' in text or '?' in text):
        if any(p in text for p in group_words):
            return ("new_topic", core)

    if '？' in text or '?' in text:
        return ("direct_question", core)

    if len(text) < 80 and any(p in text for p in directive_words):
        return ("direct_question", core)

    return ("confirm", text)


def _cap_pq(pq):
    """Cap PQ length at MAX_PQ_LENGTH at the compute_pq level."""
    if not pq or len(pq) <= MAX_PQ_LENGTH:
        return pq
    return pq[:MAX_PQ_LENGTH] + "…"


def compute_pq(mod_text, prev_last_question, topic_question):
    if not mod_text:
        return _cap_pq(prev_last_question or "")
    mtype, core = classify_mod(mod_text)
    if mtype == "true_followup":
        if prev_last_question and prev_last_question != core:
            return _cap_pq(f"{prev_last_question} → {core}")
        return _cap_pq(core)
    elif mtype == "new_topic":
        return _cap_pq(core)
    elif mtype == "name_call":
        return _cap_pq(topic_question)
    elif mtype == "direct_question":
        return _cap_pq(core)
    else:
        return _cap_pq(prev_last_question or "")


def make_respondent(speaker_id, source_file, display_name="", group_code=""):
    return {
        "speaker_id": speaker_id,
        "source_file": source_file,
        "display_name": display_name,
        "group_code": group_code,
        "profile": {"name": display_name, "age": None, "gender": "", "occupation": "", "education": ""},
        "gaming_background": {"current_games": [], "platform": [], "experience_years": None, "genre_experience": []},
        "background": {}
    }


class PQTracker:
    def __init__(self):
        self.topic_question = ""
        self.last_question = ""
        self.prev_last_question = ""

    def feed(self, mod_text):
        self.prev_last_question = self.last_question
        if not mod_text:
            return
        mtype, core = classify_mod(mod_text)
        if mtype == "new_topic":
            self.topic_question = core
            self.last_question = core
        elif mtype in ("direct_question", "true_followup"):
            self.last_question = core
        elif mtype == "name_call":
            self.last_question = self.topic_question

    def get_pq(self, mod_text):
        return compute_pq(mod_text, self.prev_last_question, self.topic_question)


# ============================================================
# group_code extraction
# ============================================================

def extract_group_code(filepath, source_file):
    """Extract group_code from file path or name."""
    fname = os.path.basename(filepath)

    # Pattern: 文字转录-海外大龄组 / 文字转录-海外年轻组
    m = re.search(r'文字转录-(.+?)(?:\.docx)?$', fname)
    if m:
        return m.group(1).strip()

    # Pattern: 座谈会笔录-G1-补访 / 座谈会笔录-G1
    m = re.search(r'[Gg](\d+(?:-补访)?)', fname)
    if m:
        return f"G{m.group(1)}"

    # Try from path
    m = re.search(r'[Gg](\d+(?:-补访)?)', source_file)
    if m:
        return f"G{m.group(1)}"

    return ""


# ============================================================
# Screening table data extraction
# ============================================================

def extract_screening_gaming_background(cells, project_type):
    """Extract gaming_background from screening table cells.
    project_type: 'deadlock' (21 cols) or 'yinghuo' (14 cols)
    """
    gb = {"current_games": [], "platform": [], "experience_years": None, "genre_experience": []}

    if project_type == 'deadlock' and len(cells) >= 21:
        # Column 6: Q1 - games played
        q1 = cells[6] if len(cells) > 6 else ""
        games = re.split(r'[\n/]+', q1)
        gb["genre_experience"] = [g.strip() for g in games if g.strip() and len(g.strip()) > 1]

        # Column 7: Q2 - Deadlock hours
        q2 = cells[7] if len(cells) > 7 else ""
        h_match = re.search(r'(\d+)\s*小时', q2)
        if h_match:
            if "deadlock_hours" not in gb:
                gb["deadlock_hours"] = int(h_match.group(1))

        # Column 11: Q4 - hours for all games
        q4 = cells[11] if len(cells) > 11 else ""
        game_hours = {}
        for line in q4.split('\n'):
            line = line.strip()
            if not line:
                continue
            # Parse: "游戏名-2000-300小时" (range), "游戏名-XXX小时", "游戏名-10年"
            gm = re.match(r'^(.+?)[-–](\d+(?:[-–]\d+)?\s*小时)', line)
            if gm:
                gname = gm.group(1).strip()
                duration = gm.group(2).strip()
                # Try to extract a single number
                h_match = re.search(r'(\d+)\s*小时', duration)
                if h_match:
                    try:
                        game_hours[gname] = int(h_match.group(1))
                    except ValueError:
                        game_hours[gname] = duration
                else:
                    game_hours[gname] = duration
                continue
            gm = re.match(r'^(.+?)[-–]\s*(\d+)\s*年', line)
            if gm:
                gname = gm.group(1).strip()
                try:
                    game_hours[gname] = int(gm.group(2)) * 365
                except ValueError:
                    pass
        if game_hours:
            gb["game_hours"] = game_hours

        # Column 12: Q5 - active games
        q5 = cells[12] if len(cells) > 12 else ""
        active = [g.strip() for g in q5.split('\n') if g.strip()]
        gb["current_games"] = active

        # Column 13: Q6 - weekly hours
        q6 = cells[13] if len(cells) > 13 else ""
        weekly = {}
        for line in q6.split('\n'):
            line = line.strip()
            if not line:
                continue
            gm = re.match(r'^(.+?)[-–]每周玩\s*(\d+.*?小时)', line)
            if gm:
                weekly[gm.group(1).strip()] = gm.group(2).strip()
        if weekly:
            gb["weekly_hours"] = weekly

        # Column 16: Q8 - OW/APEX/Valorant rank
        q8 = cells[16] if len(cells) > 16 else ""
        if q8 and q8 != '/':
            gb["shooter_ranks"] = q8.strip()

        # Column 17: Q9 - Steam account
        q9 = cells[17] if len(cells) > 17 else ""
        if q9 and q9 != '/':
            gb["platform"] = ["PC (Steam)"]
            # Extract years
            ym = re.search(r'(\d+)\s*年', q9)
            if ym:
                gb["steam_years"] = int(ym.group(1))

        # Column 19: Q11 - other MOBA/shooter games
        q11 = cells[19] if len(cells) > 19 else ""
        if q11 and q11 not in ('没有', '无', '/'):
            other_games = [g.strip() for g in q11.split('\n') if g.strip()]
            gb["other_games"] = other_games

    elif project_type == 'yinghuo' and len(cells) >= 14:
        # Column 7: Q1 - games played
        q1 = cells[7] if len(cells) > 7 else ""
        games = re.split(r'[\n/]+', q1)
        gb["genre_experience"] = [g.strip() for g in games if g.strip() and len(g.strip()) > 1]

        # Column 8: Q2 - total hours
        q2 = cells[8] if len(cells) > 8 else ""
        game_hours = {}
        for line in q2.split('\n'):
            line = line.strip()
            if not line:
                continue
            gm = re.match(r'^(.+?)[-–]\s*(\d+.*?小时|.*?月|.*?年)', line)
            if gm:
                gname = gm.group(1).strip()
                duration = gm.group(2).strip()
                game_hours[gname] = duration
        if game_hours:
            gb["game_hours"] = game_hours

        # Column 11: Q3-2 - device
        q3_2 = cells[11] if len(cells) > 11 else ""
        if q3_2:
            platforms = [p.strip() for p in q3_2.split('\n') if p.strip()]
            gb["platform"] = platforms

        # Column 12: Q3-3 - rank
        q3_3 = cells[12] if len(cells) > 12 else ""
        if q3_3 and q3_3 != '/':
            gb["rank"] = q3_3.strip()

        # Column 13: Q3-4 - maps
        q3_4 = cells[13] if len(cells) > 13 else ""
        if q3_4:
            gb["maps"] = [m.strip() for m in q3_4.split('\n') if m.strip()]

    return gb


def extract_screening_profile(cells, project_type):
    """Extract enhanced profile from screening table cells."""
    profile = {
        "name": cells[1] if len(cells) > 1 else "",
        "age": None,
        "gender": cells[2] if len(cells) > 2 else "",
        "occupation": cells[3] if len(cells) > 3 else "",
        "education": cells[5] if len(cells) > 5 else "",
    }
    # Age
    age_raw = cells[4] if len(cells) > 4 else ""
    try:
        profile["age"] = int(age_raw)
    except (ValueError, TypeError):
        profile["age"] = age_raw if age_raw else None

    # For 萤火突击: extra occupation detail in col 5
    if project_type == 'yinghuo' and len(cells) > 5:
        extra_occ = cells[5] if len(cells) > 5 else ""
        if extra_occ and extra_occ != profile["occupation"]:
            profile["occupation_detail"] = extra_occ
        profile["education"] = cells[6] if len(cells) > 6 else ""

    return profile


# ============================================================
# Post-processing: extract profile/gaming from segment text
# ============================================================

def extract_profile_from_text(text):
    """Extract profile info (age, gender, occupation) from text."""
    profile = {}

    # Age patterns
    age_patterns = [
        (r'(\d{1,2})\s*岁', 'cn'),
        (r'今年\s*(\d{1,2})', 'cn'),
        (r"(\d{1,2})\s*years?\s*old", 'en'),
        (r"I'm\s+(\d{1,2})", 'en'),
        (r"age\s*(\d{1,2})", 'en'),
        (r'年纪\s*(\d{1,2})', 'cn'),
    ]
    for pat, lang in age_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            try:
                profile["age"] = int(m.group(1))
            except ValueError:
                pass
            break

    # Gender patterns
    if re.search(r'\b男\b', text):
        profile["gender"] = "男"
    elif re.search(r'\b女\b', text):
        profile["gender"] = "女"
    elif re.search(r'\b[Mm]ale\b', text):
        profile["gender"] = "Male"
    elif re.search(r'\b[Ff]emale\b', text):
        profile["gender"] = "Female"

    # Occupation patterns (Chinese)
    occ_patterns = [
        (r'职业[是为：:]\s*(.+?)(?:[，。,\.\s]|$)', 'cn'),
        (r'我是[一]?个?\s*(.+?)(?:[，。,\.\s]|$)', 'cn'),
        (r'做(.+?)的', 'cn'),
        (r'(?:学生|工程师|程序员|设计师|运营|产品|销售|教师|医生|自由职业|公务员|个体|待业|离职|无业|金融|服务业)',
         'cn'),
        (r"I(?:'m|\s+am)\s+a?\s*(\w+(?:\s+\w+){0,3})", 'en'),
        (r"I work as a?\s*(\w+(?:\s+\w+){0,3})", 'en'),
        (r"I don't work|I don't have (?:any|a) job|I(?:'m|\s+am) unemployed", 'en_unemployed'),
        (r"I(?:'m|\s+am) (?:a|an)\s+(\w+(?:\s+\w+){0,3})", 'en'),
    ]
    for pat, lang in occ_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            if lang == 'en_unemployed':
                profile["occupation"] = "unemployed"
                break
            occ = m.group(1).strip() if m.lastindex and m.lastindex >= 1 else m.group(0).strip()
            non_occ = {'的', '人', '工作', '职业', '一个', '什么', '25 years old', 'years old',
                       'year old', 'old', 'years', 'year', 'student', 'worker', 'person',
                       'guy', 'man', 'woman', 'kid', 'player', 'gamer'}
            if len(occ) < 30 and occ.lower() not in non_occ:
                if not re.search(r'\d+\s*years?\s*old', occ, re.IGNORECASE):
                    profile["occupation"] = occ
                    break

    # Education
    edu_patterns = [
        (r'(大专|本科|硕士|博士|高中|初中|小学|研究生)', 'cn'),
        (r"(Bachelor|Master|PhD|Doctor|High\s*School|College|University)", 'en'),
    ]
    for pat, lang in edu_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            profile["education"] = m.group(1).strip()
            break

    return profile


def extract_gaming_from_text(text):
    """Extract gaming background info from text."""
    gb = {}

    # Platform
    platforms = []
    if re.search(r'\b(PC|电脑|Steam|steam|端游)\b', text):
        platforms.append("PC")
    if re.search(r'\b(手机|手游|mobile|Mobile|iOS|Android|安卓|苹果)\b', text):
        platforms.append("Mobile")
    if re.search(r'\b(主机|PS[45]|Xbox|xbox|Switch|switch|PlayStation|playstation|console|Console)\b', text):
        platforms.append("Console")
    if platforms:
        gb["platform"] = platforms

    # Game names - from 《》 brackets (Chinese) and known games list
    found_games = []
    # Extract from 《》 brackets first (most reliable for Chinese)
    bracket_games = re.findall(r'《([^》]+)》', text)
    found_games.extend(bracket_games)
    # Also check known games list
    known_games = [
        "英雄联盟", "LOL", "Dota2", "Dota", "刀塔", "守望先锋", "Overwatch", "OW",
        "APEX", "Apex", "Apex英雄", "瓦罗兰特", "Valorant", "无畏契约",
        "CS", "CSGO", "CS2", "穿越火线", "CF", "使命召唤", "COD", "Call of Duty",
        "绝地求生", "PUBG", "和平精英", "王者荣耀", "Honor of Kings",
        "Deadlock", "死锁", "漫威争锋", "Marvel Rivals", "萤火突击",
        "Fortnite", "堡垒之夜", "R6", "彩虹六号", "Rainbow Six",
        "战地", "Battlefield", "命运2", "Destiny 2", "Warframe",
        "泰坦陨落", "Titanfall", "最终幻想", "FFXIV", "FF14",
        "Rust", "DayZ", "ARK", "War Thunder", "战争雷霆",
        "World of Tanks", "坦克世界", "逃离塔科夫", "EFT", "Escape from Tarkov",
        "暗区突围", "Arena Breakout", "三角洲行动", "Delta Force",
        "THE FINALS", "Finals", "Helldivers", "绝地潜兵",
    ]
    for game in known_games:
        if game.lower() in text.lower():
            found_games.append(game)
    if found_games:
        gb["current_games"] = list(set(found_games))[:10]

    # Experience years
    exp_patterns = [
        (r'玩了\s*(\d+)\s*年', 'cn'),
        (r'(\d+)\s*年[以之]?前', 'cn'),
        (r'(\d+)\s*年', 'cn'),
        (r'playing\s+(?:for\s+)?(\d+)\s*years?', 'en'),
        (r'(\d+)\s*years?\s*(?:of\s+)?(?:gaming|playing|experience)', 'en'),
        (r'since\s*(\d{4})', 'en'),
    ]
    for pat, lang in exp_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            try:
                val = int(m.group(1))
                if lang == 'en' and 'since' in pat and val >= 2000:
                    val = 2026 - val
                if 1 <= val <= 50:
                    gb["experience_years"] = val
            except ValueError:
                pass
            break

    # Hours
    h_match = re.search(r'(\d+)\s*(?:小时|hours?|hrs?|h)', text, re.IGNORECASE)
    if h_match:
        try:
            gb["hours_mentioned"] = int(h_match.group(1))
        except ValueError:
            pass

    return gb


# ============================================================
# Post-processing: name validation helpers
# ============================================================

def _is_valid_name(name):
    """Check if a name looks like a real person's name."""
    if not name or len(name) < 2:
        return False
    if re.match(r'^P\d{3}$', name):
        return False
    if re.match(r'^SPEAKER_\d+$', name) or 'STF#' in name:
        return False
    if len(name) == 1:
        return False
    if re.search(r'[一-鿿]', name):
        bad_words = ['喜欢', '游戏', '打', '玩', '的', '人', '了', '是', '我', '你', '他', '她',
                      '什么', '怎么', '因为', '所以', '但是', '如果', '可以', '没有',
                      '主播', '直播', '视频', '抖音', 'B站', '朋友', '同学', '同事',
                      '参加', '报名', '访谈', '访问', '调研', '测试', '研究']
        for word in bad_words:
            if word in name:
                return False
        cn_chars = re.findall(r'[一-鿿]', name)
        if len(cn_chars) > 6:
            return False
    if re.match(r'^[A-Z]', name):
        return True
    if re.match(r'^[一-鿿]{2,4}$', name):
        return True
    return True


def _find_better_name(segments, speaker_id):
    """Try to find a better display name from segment text."""
    for s in segments:
        if s["speaker_id"] != speaker_id:
            continue
        text = s["original_text"]
        # Try "I'm [Name]" or "My name is [Name]"
        m = re.search(r"(?:I'm|I am|my name is|name's)\s+([A-Z][a-z]+(?:\s+[A-Z]\.?)?)", text)
        if m:
            name = m.group(1).strip()
            if _is_valid_name(name):
                return name
        # Try "[Name]. I'm..." or "[Name], I'm..." at start of text
        m = re.match(r"^([A-Z][a-z]+(?:\s+[A-Z]\.?)?)[,.]\s+(?:I(?:'m|\s+am)|my name)", text)
        if m:
            name = m.group(1).strip()
            if _is_valid_name(name):
                return name
        # Try "我叫 [Name]" or "我是 [Name]"
        m = re.search(r'(?:我是|我叫|叫)\s*([一-鿿]{2,4})\s*(?:[，。,\.\s]|$)', text)
        if m:
            name = m.group(1).strip()
            if _is_valid_name(name):
                return name
        # Try "大家好，我是XXX"
        m = re.search(r'大家好[,，]\s*(?:我是|我叫)\s*([一-鿿]{2,4})', text)
        if m:
            name = m.group(1).strip()
            if _is_valid_name(name):
                return name
        # Try "大家好，我叫XXX，今年XX岁" (Chinese focus group intro)
        m = re.search(r'大家好[,，]\s*(?:我是|我叫)\s*([一-鿿A-Za-z0-9_-]{2,20})[,，]', text)
        if m:
            name = m.group(1).strip()
            if _is_valid_name(name):
                return name
    return None


# ============================================================
# Post-processing pipeline
# ============================================================

def post_process(respondents, segments, source_file):
    """Apply all post-processing fixes to respondents and segments."""

    # 1. Fill empty display_names
    for r in respondents:
        if not r["display_name"]:
            # Try to extract from segment text
            best_name = _find_better_name(segments, r["speaker_id"])
            if best_name:
                r["display_name"] = best_name
            else:
                # Fallback: use speaker_id
                r["display_name"] = r["speaker_id"]

    # 2. Extract profile from segment text for respondents without profile data
    for r in respondents:
        p = r["profile"]
        if not p.get("name") or p.get("name") == r.get("display_name", ""):
            p["name"] = r["display_name"]

        needs_profile = (p.get("age") is None and not p.get("gender") and not p.get("occupation"))
        if needs_profile:
            for s in segments[:10]:  # Check first 10 segments
                if s["speaker_id"] == r["speaker_id"]:
                    extracted = extract_profile_from_text(s["original_text"])
                    for k, v in extracted.items():
                        if v and not p.get(k):
                            p[k] = v

    # 3. Extract gaming background from segment text for respondents without GB data
    for r in respondents:
        gb = r["gaming_background"]
        needs_gb = (not gb.get("current_games") and not gb.get("platform") and gb.get("experience_years") is None)
        if needs_gb:
            for s in segments[:10]:
                if s["speaker_id"] == r["speaker_id"]:
                    extracted = extract_gaming_from_text(s["original_text"])
                    for k, v in extracted.items():
                        if v and not gb.get(k):
                            gb[k] = v

    # 4. Backfill empty preceding_questions
    # Only backfill from a real question (contains ? or ？), not casual banter
    first_real_pq = ""
    for s in segments:
        pq = s["preceding_question"]
        if pq and ('?' in pq or '？' in pq or EN_QUESTION_RE.search(pq)):
            first_real_pq = pq
            break

    # For segments at the start of the file with empty PQ, use first_real_pq or "开场介绍"
    for i, s in enumerate(segments):
        if not s["preceding_question"]:
            if first_real_pq and i < 10:
                s["preceding_question"] = first_real_pq
            elif i < 10:
                s["preceding_question"] = "开场介绍"
            else:
                s["preceding_question"] = "[无前置提问]"

    # 5. Merge ultra-short consecutive same-speaker segments (≤5 chars)
    merged_segments = []
    for s in segments:
        text = s["original_text"].strip()
        if merged_segments and merged_segments[-1]["speaker_id"] == s["speaker_id"]:
            if len(text) <= 5:
                # Merge into previous
                merged_segments[-1]["original_text"] += " " + text
                merged_segments[-1]["char_count"] = len(merged_segments[-1]["original_text"])
                continue
            elif len(merged_segments[-1]["original_text"].strip()) <= 5:
                # Previous was short, merge this into previous
                prev_text = merged_segments[-1]["original_text"]
                merged_segments[-1]["original_text"] = prev_text + " " + text
                merged_segments[-1]["char_count"] = len(merged_segments[-1]["original_text"])
                continue
        merged_segments.append(s)

    # 6. Ensure display_name is set in profile, and validate name
    for r in respondents:
        # Validate display_name: reject names that are clearly not real names
        dn = r.get("display_name", "")
        if dn and not _is_valid_name(dn):
            # Try to find a better name from segments
            better_name = _find_better_name(segments, r["speaker_id"])
            if better_name:
                r["display_name"] = better_name
            else:
                r["display_name"] = r["speaker_id"]

        if not r["profile"].get("name") or not _is_valid_name(r["profile"].get("name", "")):
            r["profile"]["name"] = r["display_name"]

    return respondents, merged_segments


# ============================================================
# PARAGRAPH PARSER
# ============================================================

SPEAKER_PATTERNS = [
    re.compile(r'^(SPEAKER_\d+)\([^)]*\):\s*(.*)'),
    re.compile(r'^(\d+-G\d+-[A-Za-z0-9_]+)\([^)]*\):\s*(.*)'),
    re.compile(r'^([MP]):\s*(.*)'),
    re.compile(r'^([MP])：\s*(.*)'),
    re.compile(r'^([A-Z][A-Za-z.\' -]+(?:[A-Z]\.[A-Z]?)?):\s*(.*)'),
    re.compile(r'^(主持人)\([^)]*\):\s*(.*)'),
    re.compile(r'^(STF#说话人\d+)\([^)]*\):\s*(.*)'),
    re.compile(r'^([男女])：\s*(.*)'),
    re.compile(r'^([A-Za-z0-9_ -]+)\([^)]*\):\s*(.*)'),
]


def parse_paragraph_file(filepath, source_file, moderator_names=None, is_text_transcript=False):
    if moderator_names is None:
        moderator_names = {"主持人", "Moderator", "moderator", "M", "Morgan P.", "SPEAKER_05"}

    doc = Document(filepath)
    labeled_paragraphs = []
    i = 0
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    while i < len(paragraphs):
        text = paragraphs[i]
        if not text:
            i += 1
            continue

        matched = False
        for pat in SPEAKER_PATTERNS:
            m = pat.match(text)
            if m:
                speaker = m.group(1).strip()
                content = m.group(2).strip() if m.lastindex and m.lastindex >= 2 else ""

                if not content and i + 1 < len(paragraphs):
                    next_text = paragraphs[i + 1]
                    is_new_speaker = False
                    for pat2 in SPEAKER_PATTERNS:
                        if pat2.match(next_text):
                            is_new_speaker = True
                            break
                    if not is_new_speaker:
                        content = next_text
                        i += 1

                if content:
                    if content in ("主持人", "Moderator", "moderator") or content.startswith("参与者") or content.startswith("Participant"):
                        matched = True
                        break
                    labeled_paragraphs.append({"speaker": speaker, "content": content})
                matched = True
                break
        i += 1

    if not labeled_paragraphs:
        return [], []

    # Identify moderator
    speaker_counts = Counter(p["speaker"] for p in labeled_paragraphs)
    moderator = None
    for name in moderator_names:
        if name in speaker_counts:
            moderator = name
            break
    if not moderator:
        moderator = speaker_counts.most_common(1)[0][0]

    if "M" in speaker_counts and moderator != "M":
        if speaker_counts.get("M", 0) > 5:
            moderator = "M"

    # Build participants
    moderator_labels = {"M", "Moderator", "moderator", "主持人", "男", "女", "SPEAKER_05", "STF#说话人01"}
    participants = {}
    seq = 0
    for sp in speaker_counts:
        if sp == moderator or sp in moderator_labels:
            continue
        seq += 1
        sid = f"P{seq:03d}"
        participants[sp] = sid

    # Build transcript rows
    rows = []
    tracker = PQTracker()
    for p in labeled_paragraphs:
        sp = p["speaker"]
        content = p["content"]
        if sp == moderator:
            tracker.feed(content)
            # For text transcription files, only track real questions as mod context
            if is_text_transcript:
                mtype, _ = classify_mod(content)
                if mtype in ("direct_question", "new_topic"):
                    rows.append({"mod": content, "user": "", "sid": None})
                else:
                    # Keep row but mark as non-question context
                    rows.append({"mod": "", "user": "", "sid": None})
            else:
                rows.append({"mod": content, "user": "", "sid": None})
        elif sp in participants:
            sid = participants[sp]
            pq = tracker.get_pq(rows[-1]["mod"] if rows and rows[-1]["mod"] else "")
            rows.append({"mod": "", "user": content, "sid": sid, "pq": pq})

    # Build segments
    segments = []
    cur = None
    for r in rows:
        if not r["user"]:
            if cur:
                cur["original_text"] = " ".join(cur["text_parts"])
                segments.append(cur)
                cur = None
            continue
        if cur is None:
            cur = {"speaker_id": r["sid"], "speaker_role": "interviewee",
                   "preceding_question": r.get("pq", ""), "text_parts": [r["user"]]}
        elif cur["speaker_id"] == r["sid"]:
            cur["text_parts"].append(r["user"])
        else:
            cur["original_text"] = " ".join(cur["text_parts"])
            segments.append(cur)
            cur = {"speaker_id": r["sid"], "speaker_role": "interviewee",
                   "preceding_question": r.get("pq", ""), "text_parts": [r["user"]]}
    if cur:
        cur["original_text"] = " ".join(cur["text_parts"])
        segments.append(cur)

    group_code = extract_group_code(filepath, source_file)
    respondents = [make_respondent(sid, source_file, sp, group_code) for sp, sid in participants.items()]
    return respondents, segments


# ============================================================
# TABLE PARSERS
# ============================================================

def parse_table_focus_group(filepath, source_file):
    """Parse focus group with 2 tables: Table[0]=screening, Table[1]=transcript."""
    doc = Document(filepath)
    if len(doc.tables) < 2:
        return None, []

    group_code = extract_group_code(filepath, source_file)

    # Determine project type from path
    project_type = None
    if "Deadlock" in source_file:
        project_type = 'deadlock'
    elif "萤火突击" in source_file:
        project_type = 'yinghuo'

    # Table 0: Screening → participants
    profile_table = doc.tables[0]
    participants = {}
    for row in profile_table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 2:
            continue
        seq = cells[0]
        if not seq.isdigit():
            continue
        display_name = cells[1] if len(cells) > 1 else ""

        if project_type:
            profile = extract_screening_profile(cells, project_type)
            gb = extract_screening_gaming_background(cells, project_type)
        else:
            profile = {"name": display_name, "age": None, "gender": cells[2] if len(cells) > 2 else "",
                       "occupation": cells[3] if len(cells) > 3 else "", "education": cells[5] if len(cells) > 5 else ""}
            try:
                profile["age"] = int(cells[4]) if len(cells) > 4 and cells[4] else None
            except (ValueError, TypeError):
                profile["age"] = cells[4] if len(cells) > 4 and cells[4] else None
            gb = {"current_games": [], "platform": [], "experience_years": None, "genre_experience": []}

        pid = f"P{seq.zfill(3)}"
        participants[seq] = {
            "speaker_id": pid,
            "source_file": source_file,
            "display_name": display_name,
            "group_code": group_code,
            "profile": profile,
            "gaming_background": gb,
            "background": {}
        }

    # Table 1: Transcript
    transcript_table = doc.tables[1]
    rows = []
    tracker = PQTracker()
    speaker_pattern = re.compile(r'^(\d+|[多其]数人|其他人)[：:]')

    for row in transcript_table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        mod_raw = cells[0] if len(cells) > 0 else ""
        user_raw = cells[1] if len(cells) > 1 else ""

        mod_text = ""
        if mod_raw:
            if mod_raw.startswith('M：'): mod_text = mod_raw[2:]
            elif mod_raw.startswith('M:'): mod_text = mod_raw[2:]
            else: mod_text = mod_raw

        content = ""
        speaker_id = None
        if user_raw:
            m = speaker_pattern.match(user_raw)
            if m:
                speaker_marker = m.group(1)
                content = user_raw[m.end():]
                if speaker_marker.isdigit() and speaker_marker in participants:
                    speaker_id = participants[speaker_marker]["speaker_id"]

        if mod_text:
            tracker.feed(mod_text)

        if speaker_id and content:
            pq = tracker.get_pq(mod_text)
            rows.append({"speaker_id": speaker_id, "content": content, "preceding_question": pq})

    # Merge consecutive same-speaker
    segments = []
    cur = None
    for r in rows:
        if cur is None:
            cur = {"speaker_id": r["speaker_id"], "speaker_role": "interviewee",
                   "preceding_question": r["preceding_question"], "text_parts": [r["content"]]}
        elif cur["speaker_id"] == r["speaker_id"]:
            cur["text_parts"].append(r["content"])
        else:
            cur["original_text"] = " ".join(cur["text_parts"])
            segments.append(cur)
            cur = {"speaker_id": r["speaker_id"], "speaker_role": "interviewee",
                   "preceding_question": r["preceding_question"], "text_parts": [r["content"]]}
    if cur:
        cur["original_text"] = " ".join(cur["text_parts"])
        segments.append(cur)

    return list(participants.values()), segments


def parse_table_single_user(filepath, source_file):
    """Parse 1-on-1 interview table: 1 table, 2 cols. NO merging."""
    doc = Document(filepath)
    if not doc.tables:
        return None, []

    group_code = extract_group_code(filepath, source_file)
    table = doc.tables[0]
    participant_id = "P001"
    rows = []
    tracker = PQTracker()

    first_row = True
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 2:
            continue
        if first_row:
            if cells[0] == '主持人' and cells[1] == '用户':
                first_row = False
                continue
            first_row = False

        mod_raw = cells[0]
        user_raw = cells[1]

        mod_text = ""
        if mod_raw:
            if mod_raw.startswith('M：'): mod_text = mod_raw[2:]
            elif mod_raw.startswith('M:'): mod_text = mod_raw[2:]
            else: mod_text = mod_raw

        content = ""
        if user_raw:
            m = re.match(r'^\d+\s*[：:]', user_raw)
            if m:
                content = user_raw[m.end():]
            else:
                content = user_raw

        if mod_text or content:
            tracker.feed(mod_text)
            pq = tracker.get_pq(mod_text)
            if content:
                rows.append({"content": content, "preceding_question": pq})

    segments = []
    for j, r in enumerate(rows):
        segments.append({
            "speaker_id": participant_id,
            "speaker_role": "interviewee",
            "preceding_question": r["preceding_question"],
            "original_text": r["content"],
        })

    respondent = make_respondent(participant_id, source_file, "", group_code)
    return [respondent], segments


def parse_table_focus_group_1table(filepath, source_file):
    """Parse focus group with 1 table (no screening): names in user column."""
    doc = Document(filepath)
    if not doc.tables:
        return None, []

    group_code = extract_group_code(filepath, source_file)
    table = doc.tables[0]
    participants = {}
    rows = []
    tracker = PQTracker()

    name_patterns = [
        re.compile(r'^(G\d+-[A-Za-z0-9_]+)[：:]\s*'),
        re.compile(r'^(\d+-[^：:]+)[：:]\s*'),
        re.compile(r'^(\d+)[：:]\s*'),
        re.compile(r'^([一-鿿]{2,4})[：:]\s*'),
    ]

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 2:
            continue
        mod_raw = cells[0]
        user_raw = cells[1]

        mod_text = ""
        if mod_raw:
            if mod_raw.startswith('M：'): mod_text = mod_raw[2:]
            elif mod_raw.startswith('M:'): mod_text = mod_raw[2:]
            else: mod_text = mod_raw

        content = ""
        speaker_id = None
        if user_raw:
            for pat in name_patterns:
                m = pat.match(user_raw)
                if m:
                    name = m.group(1)
                    content = user_raw[m.end():]
                    if name not in participants:
                        seq = len(participants) + 1
                        pid = f"P{seq:03d}"
                        participants[name] = {
                            "speaker_id": pid,
                            "source_file": source_file,
                            "display_name": name,
                            "group_code": group_code,
                            "profile": {"name": name, "age": None, "gender": "", "occupation": "", "education": ""},
                            "gaming_background": {"current_games": [], "platform": [], "experience_years": None, "genre_experience": []},
                            "background": {}
                        }
                    speaker_id = participants[name]["speaker_id"]
                    break
            if not speaker_id:
                content = user_raw

        if mod_text:
            tracker.feed(mod_text)

        if speaker_id and content:
            pq = tracker.get_pq(mod_text)
            rows.append({"speaker_id": speaker_id, "content": content, "preceding_question": pq})

    segments = []
    cur = None
    for r in rows:
        if cur is None:
            cur = {"speaker_id": r["speaker_id"], "speaker_role": "interviewee",
                   "preceding_question": r["preceding_question"], "text_parts": [r["content"]]}
        elif cur["speaker_id"] == r["speaker_id"]:
            cur["text_parts"].append(r["content"])
        else:
            cur["original_text"] = " ".join(cur["text_parts"])
            segments.append(cur)
            cur = {"speaker_id": r["speaker_id"], "speaker_role": "interviewee",
                   "preceding_question": r["preceding_question"], "text_parts": [r["content"]]}
    if cur:
        cur["original_text"] = " ".join(cur["text_parts"])
        segments.append(cur)

    return list(participants.values()), segments


def parse_table_player_ability(filepath, source_file):
    """Parse 玩家能力: 主持人：/ N-G1-XXX： format."""
    doc = Document(filepath)
    if not doc.tables:
        return None, []

    group_code = extract_group_code(filepath, source_file)
    table = doc.tables[0]
    participants = {}
    rows = []
    tracker = PQTracker()
    name_pattern = re.compile(r'^(\d+)-([^：:]+)[：:]')

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 2:
            continue
        mod_raw = cells[0]
        user_raw = cells[1]

        mod_text = ""
        if mod_raw:
            if mod_raw.startswith('主持人：'): mod_text = mod_raw[4:]
            elif mod_raw.startswith('主持人:'): mod_text = mod_raw[4:]
            else: mod_text = mod_raw

        content = ""
        speaker_id = None
        if user_raw:
            m = name_pattern.match(user_raw)
            if m:
                seq = m.group(1)
                display_name = m.group(2).strip()
                content = user_raw[m.end():]
                if seq not in participants:
                    pid = f"P{seq.zfill(3)}"
                    participants[seq] = {
                        "speaker_id": pid,
                        "source_file": source_file,
                        "display_name": display_name,
                        "group_code": group_code,
                        "profile": {"name": display_name, "age": None, "gender": "", "occupation": "", "education": ""},
                        "gaming_background": {"current_games": [], "platform": [], "experience_years": None, "genre_experience": []},
                        "background": {}
                    }
                speaker_id = participants[seq]["speaker_id"]
            else:
                content = user_raw

        if mod_text:
            tracker.feed(mod_text)

        if speaker_id and content:
            pq = tracker.get_pq(mod_text)
            rows.append({"speaker_id": speaker_id, "content": content, "preceding_question": pq})

    segments = []
    cur = None
    for r in rows:
        if cur is None:
            cur = {"speaker_id": r["speaker_id"], "speaker_role": "interviewee",
                   "preceding_question": r["preceding_question"], "text_parts": [r["content"]]}
        elif cur["speaker_id"] == r["speaker_id"]:
            cur["text_parts"].append(r["content"])
        else:
            cur["original_text"] = " ".join(cur["text_parts"])
            segments.append(cur)
            cur = {"speaker_id": r["speaker_id"], "speaker_role": "interviewee",
                   "preceding_question": r["preceding_question"], "text_parts": [r["content"]]}
    if cur:
        cur["original_text"] = " ".join(cur["text_parts"])
        segments.append(cur)

    return list(participants.values()), segments


# ============================================================
# WEBVTT PARSER (rewritten with proper PQ extraction)
# ============================================================

def is_moderator_webvtt(text):
    """Determine if a WEBVTT text line is from the moderator."""
    text = text.strip()
    if not text:
        return False

    # Strong moderator signals
    # 1. Tag question "right?"
    if text.rstrip().endswith("right?"):
        return True

    # 2. Ends with "?" (not "right?")
    if text.rstrip().endswith("?"):
        return True

    # 3. Discourse markers (only if accompanied by facilitation/question language)
    discourse = r'^(Great|Perfect|Cool|Interesting|All right|Alright|Okay|OK|Yeah|Yep|Right|Good|Awesome|Wonderful|Fantastic|Got it|I see|Fair enough|Good to know|So|And|But|Also|Like|Um|Hmm)\.?\s'
    if re.match(discourse, text, re.IGNORECASE):
        # Short responses can be participant backchannels
        if len(text) < 30:
            return False
        # For common sentence starters (So, And, But, Also, Like), require
        # additional facilitation/question signals to avoid participant speech
        common_starters = r'^(So|And|But|Also|Like)\b'
        if re.match(common_starters, text, re.IGNORECASE):
            # Must contain a question mark or facilitation/question keywords
            if '?' in text:
                return True
            # Check for common facilitation/question patterns inline
            mod_keywords = [
                "let's", "we'll", "we're gonna", "I want to", "I'd like to",
                "what do you", "anyone else", "any questions",
                "we're going to", "can you", "could you", "tell me",
                "what about", "how about", "what else", "anything else",
                "go ahead", "go on", "continue", "tell me more",
                "we'll start", "we'll go to", "let's talk",
                "introduce", "describe", "share", "explain",
                "thank you", "thanks", "appreciate",
                "you guys can", "you guys see", "let me know",
                "before we get started", "moving on", "next up",
                "focus on", "we're gonna do", "we're gonna talk",
                "let's go", "let's move", "let's start", "let's begin",
                "I'm going to", "I'm gonna", "I will", "I'll just",
                "I just want", "I want to hear", "we want to know",
                "can you guys", "do you guys", "have you guys",
                "what do you", "how do you", "why do you",
                "what did you", "how did you", "what made you",
                "what was", "what is", "what are", "how is", "how are",
                "did you", "do you", "have you", "are you",
                "is there", "are there", "was there", "were there",
                "any other", "anyone have", "anybody have",
                "who else", "what else",
                "so we have", "so we're", "so let's", "so I'll",
                "so what", "so how", "so who", "so when",
                "I'm curious", "I'm wondering", "I wonder",
                "let's hear", "let me ask", "let me just",
                "welcome", "thank you all", "thanks everyone",
                "hello everyone", "hi everyone", "hey everyone",
                "good morning", "good afternoon", "good evening",
                "my name is", "I'm your", "I am your",
                "today we're", "we're here", "we're meeting",
                "this meeting", "this session", "this focus group",
                "this interview", "this discussion", "this study",
                "we're recording", "this recording",
                "internal purposes", "keep it",
                "don't worry", "no wrong", "no right",
                "honest", "be honest",
                "your opinion", "your experience", "your thoughts",
                "your feelings", "your perspective", "your view",
                "your feedback", "your input", "your insights",
                "your take", "what you think", "how you feel",
                "what you like", "what you don't", "what you love",
                "what you hate", "what stands out", "what stuck",
                "what's your", "what is your", "what are your",
                "how would you", "how do you", "how did you",
                "can you describe", "can you tell", "can you explain",
                "can you talk", "can you share", "can you walk",
                "could you describe", "could you tell", "could you explain",
                "could you talk", "could you share",
            ]
            if any(marker in text.lower() for marker in mod_keywords):
                return True
            return False
        return True

    # 4. Facilitation language
    facilitation = [
        "let's", "we'll", "we're gonna", "I want to", "I'd like to",
        "what do you guys", "anyone else", "any questions",
        "we're going to", "can you", "could you", "tell me",
        "what about", "how about", "what else", "anything else",
        "go ahead", "go on", "continue", "tell me more",
        "we'll start with", "we'll go to", "let's talk about",
        "introduce", "describe", "share", "explain",
        "thank you", "thanks", "appreciate",
        "we're gonna talk", "I'm gonna share",
        "you guys can", "you guys see", "let me know",
        "before we get started", "moving on", "next up",
        "we'll talk about", "we'll discuss", "we're gonna do",
        "focus on", "what we're gonna do", "what we're doing",
        "let's go to", "let's go ahead", "let's move on",
        "let's start", "let's begin", "let's get started",
        "I'm going to", "I'm gonna", "I will", "I'll just",
        "I just want", "I want to hear", "we want to know",
        "can you guys", "do you guys", "have you guys",
        "what do you", "how do you", "why do you",
        "what did you", "how did you", "why did you",
        "what made you", "how did that", "what was",
        "what is", "what are", "how is", "how are",
        "did you", "do you", "have you", "are you",
        "is there", "are there", "was there", "were there",
        "any other", "anyone have", "anybody have",
        "who else", "what else", "anything else",
        "so we have", "so we're", "so let's", "so I'll",
        "so what", "so how", "so who", "so when",
        "and then", "and so", "and what", "and how",
        "we'll circle", "we'll come", "we'll get",
        "we'll talk", "we'll go", "we'll start",
        "we'll do", "we'll see", "we'll hear",
        "I'm curious", "I'm wondering", "I wonder",
        "I'd love", "I would love", "I'd like",
        "I want to", "I wanna", "I need to",
        "let's hear", "let me ask", "let me just",
        "let me share", "let me pull", "I'm gonna share",
        "I'm going to share", "we're just", "we're almost",
        "give it", "hang on", "hold on", "one second",
        "bear with", "excuse", "pardon", "sorry",
        "welcome", "thank you all", "thanks everyone",
        "hello everyone", "hi everyone", "hey everyone",
        "good morning", "good afternoon", "good evening",
        "my name is", "I'm your", "I am your",
        "today we're", "we're here", "we're meeting",
        "this meeting", "this session", "this focus group",
        "this interview", "this discussion", "this study",
        "this research", "we're recording", "this recording",
        "the recording", "internal purposes", "keep it",
        "keep everything", "don't go around", "keep this",
        "you get paid", "I get paid", "we're all happy",
        "any questions", "before we", "last thing",
        "the last thing", "one more thing", "quick reminder",
        "just a reminder", "just to remind", "as a reminder",
        "NDA", "confidential", "privacy", "anonymous",
        "don't worry", "no wrong", "no right", "no correct",
        "honest", "truth", "the truth", "be honest",
        "your opinion", "your experience", "your thoughts",
        "your feelings", "your perspective", "your view",
        "your feedback", "your input", "your insights",
        "your take", "what you think", "how you feel",
        "what you like", "what you don't", "what you love",
        "what you hate", "what you enjoy", "what you prefer",
        "what stands out", "what stuck", "what caught",
        "what attracted", "what convinced", "what made",
        "what got you", "what drew", "what pulled",
        "what kept", "what keeps", "what makes",
        "what's your", "what is your", "what are your",
        "how would you", "how do you", "how did you",
        "how has", "how was", "how is",
        "can you describe", "can you tell", "can you explain",
        "can you talk", "can you share", "can you walk",
        "can you give", "can you provide", "can you expand",
        "could you describe", "could you tell", "could you explain",
        "could you talk", "could you share", "could you walk",
        "could you give", "could you provide", "could you expand",
        "would you say", "would you describe", "would you consider",
        "would you recommend", "would you agree",
        "tell me about", "tell me more", "tell us about",
        "tell us more", "tell me what", "tell us what",
        "talk to me", "talk to us", "talk about",
        "describe your", "describe the", "describe how",
        "explain your", "explain the", "explain how",
        "share your", "share what", "share how",
        "walk me through", "walk us through",
        "give me", "give us", "give an example",
        "what do you mean", "what does that mean",
        "how does that", "why does that", "when does that",
        "where does that", "who does that",
        "is that", "was that", "are those", "were those",
        "does that", "did that", "do those",
        "has that", "had that", "have those",
        "will that", "would that", "could that",
        "should that", "might that", "can that",
        "what about", "how about", "why about",
        "who about", "where about", "when about",
        "what else", "who else", "where else",
        "how else", "why else", "when else",
        "anything else", "anyone else", "anybody else",
        "somebody else", "something else", "someone else",
        "any other", "any more", "anything more",
        "anyone more", "anybody more", "something more",
        "someone more", "somebody more",
        "what do you think", "what do you guys think",
        "what do you all think", "what does everyone think",
        "how do you feel", "how do you guys feel",
        "what's your take", "what's everyone's take",
        "what's your opinion", "what's everyone's opinion",
        "what's your perspective", "what's everyone's perspective",
        "what's your experience", "what's everyone's experience",
        "what's your thought", "what's everyone's thought",
        "what's your impression", "what's everyone's impression",
        "what's your reaction", "what's everyone's reaction",
        "what's your sense", "what's everyone's sense",
        "what's your feeling", "what's everyone's feeling",
        "what's your view", "what's everyone's view",
        "what stands out to you", "what stood out to you",
        "what stuck with you", "what caught your attention",
        "what caught your eye", "what grabbed you",
        "what attracted you", "what drew you in",
        "what pulled you in", "what got you interested",
        "what got you excited", "what made you want",
        "what made you decide", "what made you choose",
        "what made you pick", "what made you buy",
        "what made you play", "what made you try",
        "what made you start", "what made you continue",
        "what made you keep", "what made you stay",
        "what made you leave", "what made you quit",
        "what made you stop", "what made you switch",
        "what made you change", "what made you move",
        "what keeps you", "what keeps you playing",
        "what keeps you coming", "what keeps you going",
        "what keeps you engaged", "what keeps you interested",
        "what keeps you hooked", "what keeps you invested",
        "what keeps you entertained", "what keeps you excited",
        "what keeps you motivated", "what keeps you coming back",
        "what do you like", "what do you love",
        "what do you enjoy", "what do you prefer",
        "what do you dislike", "what do you hate",
        "what do you find", "what do you consider",
        "what do you feel", "what do you think",
        "what do you believe", "what do you expect",
        "what do you want", "what do you need",
        "what do you look for", "what do you value",
        "what do you care about", "what do you prioritize",
        "what do you focus on", "what do you pay attention to",
        "what do you notice", "what do you observe",
        "what do you see", "what do you hear",
        "how do you approach", "how do you handle",
        "how do you deal", "how do you manage",
        "how do you decide", "how do you choose",
        "how do you pick", "how do you select",
        "how do you evaluate", "how do you judge",
        "how do you compare", "how do you rate",
        "how do you rank", "how do you prioritize",
        "how do you balance", "how do you allocate",
        "how do you spend", "how do you use",
        "how do you play", "how do you game",
        "how do you approach gaming", "how do you play games",
        "when did you", "when do you", "when would you",
        "where did you", "where do you", "where would you",
        "who did you", "who do you", "who would you",
        "why did you", "why do you", "why would you",
        "which did you", "which do you", "which would you",
        "how long", "how often", "how many", "how much",
        "how far", "how close", "how near", "how soon",
        "how quickly", "how fast", "how slow", "how well",
        "how good", "how bad", "how important", "how significant",
        "how different", "how similar", "how unique", "how special",
        "how common", "how rare", "how typical", "how unusual",
        "how easy", "how hard", "how difficult", "how simple",
        "how complex", "how straightforward", "how complicated",
        "how fun", "how enjoyable", "how entertaining",
        "how engaging", "how immersive", "how addictive",
        "how frustrating", "how annoying", "how disappointing",
        "how satisfying", "how rewarding", "how fulfilling",
        "how challenging", "how demanding", "how punishing",
        "how forgiving", "how accessible", "how approachable",
        "how welcoming", "how inviting", "how intimidating",
        "how scary", "how overwhelming", "how confusing",
        "how clear", "how obvious", "how intuitive",
        "how natural", "how smooth", "how polished",
        "how rough", "how janky", "how buggy", "how broken",
        "how stable", "how reliable", "how consistent",
        "how balanced", "how fair", "how competitive",
        "how casual", "how hardcore", "how serious",
        "compared to", "in comparison", "relative to",
        "versus", "vs", "against", "between",
        "what's the difference", "how do they compare",
        "how does it compare", "how does this compare",
        "which is better", "which do you prefer",
        "which one", "which game", "which version",
        "which platform", "which system", "which device",
        "which mode", "which character", "which class",
        "which weapon", "which loadout", "which build",
        "which strategy", "which approach", "which style",
        "you mentioned", "you said", "you talked about",
        "you brought up", "you referenced", "you alluded to",
        "you touched on", "you hinted at", "you suggested",
        "you implied", "you indicated", "you pointed out",
        "you noted", "you observed", "you noticed",
        "you highlighted", "you emphasized", "you stressed",
        "you focused on", "you concentrated on", "you zeroed in on",
        "you honed in on", "you drilled down", "you dug into",
        "you explored", "you investigated", "you examined",
        "you looked at", "you checked out", "you tested",
        "you tried", "you experimented", "you played around",
        "you messed around", "you fooled around", "you tinkered",
        "earlier you", "before you", "previously you",
        "a moment ago", "a minute ago", "a second ago",
        "just now", "just a moment", "just a minute",
        "just a second", "a little while", "a bit ago",
        "going back to", "coming back to", "returning to",
        "circling back", "looping back", "bringing it back",
        "tying back", "connecting back", "linking back",
        "relating back", "referring back", "harkening back",
        "building on", "expanding on", "elaborating on",
        "following up", "digging deeper", "diving deeper",
        "going deeper", "delving deeper", "probing further",
        "exploring further", "investigating further",
        "examining further", "looking further", "checking further",
        "testing further", "trying further", "experimenting further",
        "pushing further", "taking further", "carrying further",
        "extending further", "developing further", "advancing further",
        "progressing further", "moving further", "going further",
        "I want to hear", "I'd love to hear", "I'm curious to hear",
        "I'm interested to hear", "I'm eager to hear", "I'm excited to hear",
        "I want to know", "I'd love to know", "I'm curious to know",
        "I'm interested to know", "I'm eager to know", "I'm excited to know",
        "I want to understand", "I'd love to understand", "I'm curious to understand",
        "I'm interested to understand", "I'm eager to understand", "I'm excited to understand",
        "I want to learn", "I'd love to learn", "I'm curious to learn",
        "I'm interested to learn", "I'm eager to learn", "I'm excited to learn",
        "I want to see", "I'd love to see", "I'm curious to see",
        "I'm interested to see", "I'm eager to see", "I'm excited to see",
        "I want to get", "I'd love to get", "I'm curious to get",
        "I'm interested to get", "I'm eager to get", "I'm excited to get",
        "I want to find", "I'd love to find", "I'm curious to find",
        "I'm interested to find", "I'm eager to find", "I'm excited to find",
        "I want to discover", "I'd love to discover", "I'm curious to discover",
        "I'm interested to discover", "I'm eager to discover", "I'm excited to discover",
        "I want to explore", "I'd love to explore", "I'm curious to explore",
        "I'm interested to explore", "I'm eager to explore", "I'm excited to explore",
        "let's talk about", "let's discuss", "let's explore", "let's examine",
        "let's look at", "let's check out", "let's go over", "let's review",
        "let's cover", "let's address", "let's tackle", "let's handle",
        "let's deal with", "let's work through", "let's go through",
        "let's walk through", "let's run through", "let's step through",
        "let's think about", "let's consider", "let's ponder",
        "let's reflect on", "let's contemplate", "let's mull over",
        "let's brainstorm", "let's ideate", "let's generate",
        "let's come up with", "let's think of", "let's imagine",
        "let's picture", "let's visualize", "let's envision",
        "let's focus on", "let's concentrate on", "let's zero in on",
        "let's hone in on", "let's drill down", "let's dig into",
        "let's dive into", "let's delve into", "let's get into",
        "let's go into", "let's move into", "let's jump into",
        "let's transition to", "let's switch to", "let's turn to",
        "let's move on to", "let's go on to", "let's proceed to",
        "let's advance to", "let's progress to", "let's continue to",
        "let's continue with", "let's proceed with", "let's go with",
        "let's move with", "let's carry on", "let's keep going",
        "let's keep moving", "let's keep pushing", "let's keep working",
        "let's keep at it", "let's keep on", "let's keep up",
        "we'll talk about", "we'll discuss", "we'll explore", "we'll examine",
        "we'll look at", "we'll check out", "we'll go over", "we'll review",
        "we'll cover", "we'll address", "we'll tackle", "we'll handle",
        "we'll deal with", "we'll work through", "we'll go through",
        "we'll walk through", "we'll run through", "we'll step through",
        "we'll think about", "we'll consider", "we'll ponder",
        "we'll reflect on", "we'll contemplate", "we'll mull over",
        "we'll brainstorm", "we'll ideate", "we'll generate",
        "we'll come up with", "we'll think of", "we'll imagine",
        "we'll picture", "we'll visualize", "we'll envision",
        "we'll focus on", "we'll concentrate on", "we'll zero in on",
        "we'll hone in on", "we'll drill down", "we'll dig into",
        "we'll dive into", "we'll delve into", "we'll get into",
        "we'll go into", "we'll move into", "we'll jump into",
        "we'll transition to", "we'll switch to", "we'll turn to",
        "we'll move on to", "we'll go on to", "we'll proceed to",
        "we'll advance to", "we'll progress to", "we'll continue to",
        "we'll continue with", "we'll proceed with", "we'll go with",
        "we'll move with", "we'll carry on", "we'll keep going",
        "we'll keep moving", "we'll keep pushing", "we'll keep working",
        "we'll keep at it", "we'll keep on", "we'll keep up",
        "we're gonna talk about", "we're gonna discuss", "we're gonna explore",
        "we're gonna examine", "we're gonna look at", "we're gonna check out",
        "we're gonna go over", "we're gonna review", "we're gonna cover",
        "we're gonna address", "we're gonna tackle", "we're gonna handle",
        "we're gonna deal with", "we're gonna work through", "we're gonna go through",
        "we're gonna walk through", "we're gonna run through", "we're gonna step through",
        "we're gonna think about", "we're gonna consider", "we're gonna ponder",
        "we're gonna reflect on", "we're gonna contemplate", "we're gonna mull over",
        "we're gonna brainstorm", "we're gonna ideate", "we're gonna generate",
        "we're gonna come up with", "we're gonna think of", "we're gonna imagine",
        "we're gonna picture", "we're gonna visualize", "we're gonna envision",
        "we're gonna focus on", "we're gonna concentrate on", "we're gonna zero in on",
        "we're gonna hone in on", "we're gonna drill down", "we're gonna dig into",
        "we're gonna dive into", "we're gonna delve into", "we're gonna get into",
        "we're gonna go into", "we're gonna move into", "we're gonna jump into",
        "we're gonna transition to", "we're gonna switch to", "we're gonna turn to",
        "we're gonna move on to", "we're gonna go on to", "we're gonna proceed to",
        "we're gonna advance to", "we're gonna progress to", "we're gonna continue to",
        "we're gonna continue with", "we're gonna proceed with", "we're gonna go with",
        "we're gonna move with", "we're gonna carry on", "we're gonna keep going",
        "we're gonna keep moving", "we're gonna keep pushing", "we're gonna keep working",
        "we're gonna keep at it", "we're gonna keep on", "we're gonna keep up",
        "I'm going to share", "I'm going to show", "I'm going to pull up",
        "I'm going to put up", "I'm going to display", "I'm going to present",
        "I'm going to demonstrate", "I'm going to illustrate", "I'm going to exhibit",
        "I'm going to reveal", "I'm going to unveil", "I'm going to expose",
        "I'm going to disclose", "I'm going to uncover", "I'm going to publish",
        "I'm going to release", "I'm going to launch", "I'm going to roll out",
        "I'm going to deploy", "I'm going to implement", "I'm going to execute",
        "I'm going to perform", "I'm going to conduct", "I'm going to carry out",
        "I'm going to undertake", "I'm going to engage in", "I'm going to participate in",
        "I'm going to take part in", "I'm going to join in", "I'm going to get involved in",
        "I'm going to be part of", "I'm going to contribute to", "I'm going to add to",
        "I'm going to supplement", "I'm going to complement", "I'm going to enhance",
        "I'm going to improve", "I'm going to upgrade", "I'm going to update",
        "I'm going to modify", "I'm going to change", "I'm going to alter",
        "I'm going to adjust", "I'm going to tweak", "I'm going to fine-tune",
        "I'm going to optimize", "I'm going to refine", "I'm going to polish",
        "I'm going to perfect", "I'm going to complete", "I'm going to finish",
        "I'm going to finalize", "I'm going to conclude", "I'm going to end",
        "I'm going to terminate", "I'm going to cease", "I'm going to stop",
        "I'm going to halt", "I'm going to pause", "I'm going to suspend",
        "I'm going to resume", "I'm going to restart", "I'm going to reboot",
        "I'm going to reload", "I'm going to refresh", "I'm going to renew",
        "I'm going to revitalize", "I'm going to rejuvenate", "I'm going to reinvigorate",
        "I'm going to reenergize", "I'm going to recharge", "I'm going to refuel",
        "I'm going to replenish", "I'm going to restock", "I'm going to resupply",
        "I'm going to refill", "I'm going to reload", "I'm going to repack",
        "I'm going to repackage", "I'm going to rebrand", "I'm going to rename",
        "I'm going to retitle", "I'm going to relabel", "I'm going to retag",
        "I'm going to recategorize", "I'm going to reclassify", "I'm going to reorganize",
        "I'm going to restructure", "I'm going to rearrange", "I'm going to reorder",
        "I'm going to resort", "I'm going to reshuffle", "I'm going to remix",
        "I'm going to recombine", "I'm going to reassemble", "I'm going to reconstruct",
        "I'm going to rebuild", "I'm going to remake", "I'm going to redo",
        "I'm going to rework", "I'm going to revise", "I'm going to rewrite",
        "I'm going to rephrase", "I'm going to reword", "I'm going to reformulate",
        "I'm going to reexpress", "I'm going to rearticulate", "I'm going to restate",
        "I'm going to reiterate", "I'm going to repeat", "I'm going to recap",
        "I'm going to recapitulate", "I'm going to summarize", "I'm going to sum up",
        "I'm going to synopsize", "I'm going to abstract", "I'm going to brief",
        "do you want to", "would you like to", "do you care to",
        "do you wish to", "do you plan to", "do you intend to",
        "do you hope to", "do you expect to", "do you anticipate",
        "do you foresee", "do you predict", "do you project",
        "do you estimate", "do you calculate", "do you compute",
        "do you figure", "do you reckon", "do you suppose",
        "do you assume", "do you presume", "do you guess",
        "do you imagine", "do you picture", "do you visualize",
        "do you envision", "do you conceive", "do you perceive",
        "do you sense", "do you feel", "do you think",
        "do you believe", "do you consider", "do you regard",
        "do you view", "do you see", "do you understand",
        "do you comprehend", "do you grasp", "do you get",
        "do you follow", "do you track", "do you monitor",
        "do you observe", "do you watch", "do you notice",
        "can you tell me", "can you share with me", "can you explain to me",
        "can you describe to me", "can you walk me through", "can you talk me through",
        "can you guide me through", "can you lead me through", "can you show me",
        "can you demonstrate", "can you illustrate", "can you exemplify",
        "can you give me", "can you provide me", "can you offer me",
        "can you present me", "can you furnish me", "can you supply me",
        "can you deliver me", "can you hand me", "can you pass me",
        "can you send me", "can you email me", "can you text me",
        "can you message me", "can you DM me", "can you PM me",
        "can you contact me", "can you reach me", "can you get in touch",
        "can you connect with me", "can you link up with me", "can you meet with me",
        "can you speak with me", "can you talk with me", "can you chat with me",
        "can you converse with me", "can you communicate with me", "can you interact with me",
        "can you engage with me", "can you collaborate with me", "can you work with me",
        "can you partner with me", "can you team up with me", "can you join me",
        "could you tell me", "could you share with me", "could you explain to me",
        "could you describe to me", "could you walk me through", "could you talk me through",
        "could you guide me through", "could you lead me through", "could you show me",
        "could you demonstrate", "could you illustrate", "could you exemplify",
        "could you give me", "could you provide me", "could you offer me",
        "could you present me", "could you furnish me", "could you supply me",
        "could you deliver me", "could you hand me", "could you pass me",
        "could you send me", "could you email me", "could you text me",
        "could you message me", "could you DM me", "could you PM me",
        "could you contact me", "could you reach me", "could you get in touch",
        "could you connect with me", "could you link up with me", "could you meet with me",
        "could you speak with me", "could you talk with me", "could you chat with me",
        "could you converse with me", "could you communicate with me", "could you interact with me",
        "could you engage with me", "could you collaborate with me", "could you work with me",
        "could you partner with me", "could you team up with me", "could you join me",
        "would you tell me", "would you share with me", "would you explain to me",
        "would you describe to me", "would you walk me through", "would you talk me through",
        "would you guide me through", "would you lead me through", "would you show me",
        "would you demonstrate", "would you illustrate", "would you exemplify",
        "would you give me", "would you provide me", "would you offer me",
        "would you present me", "would you furnish me", "would you supply me",
        "would you deliver me", "would you hand me", "would you pass me",
        "would you send me", "would you email me", "would you text me",
        "would you message me", "would you DM me", "would you PM me",
        "would you contact me", "would you reach me", "would you get in touch",
        "would you connect with me", "would you link up with me", "would you meet with me",
        "would you speak with me", "would you talk with me", "would you chat with me",
        "would you converse with me", "would you communicate with me", "would you interact with me",
        "would you engage with me", "would you collaborate with me", "would you work with me",
        "would you partner with me", "would you team up with me", "would you join me",
        "will you tell me", "will you share with me", "will you explain to me",
        "will you describe to me", "will you walk me through", "will you talk me through",
        "will you guide me through", "will you lead me through", "will you show me",
        "will you demonstrate", "will you illustrate", "will you exemplify",
        "will you give me", "will you provide me", "will you offer me",
        "will you present me", "will you furnish me", "will you supply me",
        "will you deliver me", "will you hand me", "will you pass me",
        "will you send me", "will you email me", "will you text me",
        "will you message me", "will you DM me", "will you PM me",
        "will you contact me", "will you reach me", "will you get in touch",
        "will you connect with me", "will you link up with me", "will you meet with me",
        "will you speak with me", "will you talk with me", "will you chat with me",
        "will you converse with me", "will you communicate with me", "will you interact with me",
        "will you engage with me", "will you collaborate with me", "will you work with me",
        "will you partner with me", "will you team up with me", "will you join me",
    ]

    text_lower = text.lower()
    for phrase in facilitation:
        if phrase.lower() in text_lower:
            return True

    return False


def extract_webvtt_question(text):
    """Extract the core question from a moderator text line."""
    text = text.strip()
    if len(text) <= 80:
        if len(text) > MAX_PQ_LENGTH:
            return text[:MAX_PQ_LENGTH] + "…"
        return text

    # For long moderator text, try to extract the question part
    # Look for the last sentence that contains a question mark
    parts = re.split(r'([.?!])', text)
    sentences = []
    for i in range(0, len(parts) - 1, 2):
        sent = parts[i] + parts[i+1]
        if sent.strip():
            sentences.append(sent.strip())
    if len(parts) % 2 == 1 and parts[-1].strip():
        sentences.append(parts[-1].strip())

    # Find sentences with questions
    q_sentences = [s for s in sentences if '?' in s]
    if q_sentences:
        candidate = q_sentences[-1]
        if len(candidate) > 80:
            clauses = re.split(r'[,]+', candidate)
            clauses = [c.strip() for c in clauses if c.strip() and len(c.strip()) >= 4]
            if len(clauses) > 4:
                candidate = ', '.join(clauses[-4:])
        if len(candidate) > MAX_PQ_LENGTH:
            candidate = candidate[:MAX_PQ_LENGTH] + "…"
        return candidate

    if len(sentences) >= 2:
        candidate = sentences[-1]
        if len(candidate) < 20 and len(sentences) >= 2:
            candidate = sentences[-2] + candidate
    else:
        candidate = text

    if len(candidate) > 80:
        clauses = re.split(r'[,]+', candidate)
        clauses = [c.strip() for c in clauses if c.strip() and len(c.strip()) >= 4]
        if len(clauses) > 4:
            candidate = ', '.join(clauses[-4:])

    if len(candidate) > MAX_PQ_LENGTH:
        candidate = candidate[:MAX_PQ_LENGTH] + "…"
    return candidate


def parse_webvtt(filepath, source_file):
    """Parse WEBVTT subtitle files with proper moderator question extraction."""
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    group_code = extract_group_code(filepath, source_file)
    timestamp_re = re.compile(r'^\d{2}:\d{2}:\d{2}\.\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}\.\d{3}$')

    # Extract all content lines
    text_lines = []
    for text in paragraphs:
        if not text or text == "WEBVTT" or text.isdigit():
            continue
        if timestamp_re.match(text):
            continue
        text_lines.append(text)

    if not text_lines:
        return [], []

    # Classify each line as moderator or participant
    classified = []
    for line in text_lines:
        is_mod = is_moderator_webvtt(line)
        classified.append({"text": line, "is_mod": is_mod})

    # Build segments: moderator lines become questions, participant lines become answers
    # Consecutive same-type lines are merged
    merged_lines = []
    for item in classified:
        if merged_lines and merged_lines[-1]["is_mod"] == item["is_mod"]:
            merged_lines[-1]["text"] += " " + item["text"]
        else:
            merged_lines.append(item.copy())

    # Assign PQs to participant segments
    rows = []
    current_question = ""
    topic_question = ""

    for item in merged_lines:
        if item["is_mod"]:
            q = extract_webvtt_question(item["text"])
            if q:
                current_question = q
                # Check if this is a new topic question
                if any(marker in item["text"].lower() for marker in [
                    "let's get started", "let's begin", "first of all", "we'll start",
                    "introduce yourself", "tell me about yourself", "your name",
                    "we're gonna talk about", "today we're talking", "today we're going to",
                    "let's talk about", "let's discuss", "the topic today",
                    "we're going to be discussing", "we're here to talk about",
                    "the purpose of today", "what we're doing today", "what we're gonna do",
                ]):
                    topic_question = q
        else:
            pq = current_question or topic_question
            rows.append({"content": item["text"], "preceding_question": pq})

    # Build segments
    segments = []
    for j, r in enumerate(rows):
        segments.append({
            "speaker_id": "P001",
            "speaker_role": "interviewee",
            "preceding_question": r["preceding_question"],
            "original_text": r["content"],
        })

    respondent = make_respondent("P001", source_file, "", group_code)
    return [respondent], segments


# ============================================================
# XLSX PARSER (restructured per §6.4)
# ============================================================

# Sheets that are content (contain actual respondent data)
CONTENT_SHEET_KEYWORDS = [
    "玩家行为", "乐趣整理", "经验认知", "乐趣对比",
    "记录表", "特征诉求", "玩家行为&乐趣",
    "ABIvsDF记录表", "经验认知乐趣对比",
]

# Sheets that are reference (study design, scripts, screening)
REFERENCE_SHEET_KEYWORDS = [
    "研究计划", "样本条件", "用户列表", "测试安排",
    "访谈大纲", "Manuscript", "大纲",
    "ABIvsDF新手访谈大纲",
]


def parse_xlsx(filepath, source_file):
    """Parse xlsx files per §6.4: distinguish content vs reference sheets."""
    wb = load_workbook(filepath, data_only=True)
    all_segments = []
    all_respondents = []
    group_code = extract_group_code(filepath, source_file)

    # First pass: extract respondent info from 用户列表 sheet
    respondent_map = {}
    for sheet_name in wb.sheetnames:
        if "用户列表" in sheet_name:
            ws = wb[sheet_name]
            sheet_rows = []
            for row in ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True):
                sheet_rows.append([str(c) if c is not None else "" for c in row])

            # Find the header row
            header_row_idx = None
            for i, row in enumerate(sheet_rows):
                if any("ID" in str(c) or "Name" in str(c) for c in row):
                    header_row_idx = i
                    break

            if header_row_idx is not None:
                header = sheet_rows[header_row_idx]
                for row in sheet_rows[header_row_idx + 1:]:
                    if len(row) < 2:
                        continue
                    rid = row[0].strip() if len(row) > 0 else ""
                    name = row[1].strip() if len(row) > 1 else ""
                    if not rid or rid == "None":
                        continue
                    # Extract group
                    group = ""
                    gm = re.match(r'G(\d+)', rid)
                    if gm:
                        group = f"G{gm.group(1)}"
                    elif len(row) > 3:
                        group_raw = row[3].strip() if len(row) > 3 else ""
                        gm2 = re.search(r'(\d)', group_raw)
                        if gm2:
                            group = f"G{gm2.group(1)}"

                    gender = row[5].strip() if len(row) > 5 else ""
                    age_raw = row[6].strip() if len(row) > 6 else ""
                    try:
                        age = int(age_raw) if age_raw else None
                    except ValueError:
                        age = None

                    platform = row[8].strip() if len(row) > 8 else ""

                    pid = f"P{len(respondent_map) + 1:03d}"
                    respondent_map[rid] = {
                        "speaker_id": pid,
                        "source_file": source_file,
                        "display_name": name,
                        "group_code": group or group_code,
                        "profile": {"name": name, "age": age, "gender": gender, "occupation": "", "education": ""},
                        "gaming_background": {"current_games": [], "platform": [platform] if platform else [], "experience_years": None, "genre_experience": []},
                        "background": {}
                    }

    # If no 用户列表, create respondents from content sheets
    if not respondent_map:
        # Extract from content sheets
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            if ws.max_row < 2:
                continue
            sheet_rows = []
            for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 20), values_only=True):
                sheet_rows.append([str(c) if c is not None else "" for c in row])
            if not sheet_rows:
                continue

            header = sheet_rows[0]
            # Look for player ID columns
            for ci, h in enumerate(header):
                if not h:
                    continue
                # Check if this column has player data
                for r in sheet_rows[1:]:
                    if ci < len(r) and r[ci]:
                        val = r[ci].strip()
                        # Match player ID patterns
                        if re.match(r'^[PG]\d+', val) or re.match(r'^G\d+', val):
                            if val not in respondent_map:
                                pid = f"P{len(respondent_map) + 1:03d}"
                                # Extract group
                                gm = re.match(r'G(\d+)', val)
                                grp = f"G{gm.group(1)}" if gm else group_code
                                respondent_map[val] = {
                                    "speaker_id": pid,
                                    "source_file": source_file,
                                    "display_name": val,
                                    "group_code": grp,
                                    "profile": {"name": val, "age": None, "gender": "", "occupation": "", "education": ""},
                                    "gaming_background": {"current_games": [], "platform": [], "experience_years": None, "genre_experience": []},
                                    "background": {}
                                }

    all_respondents = list(respondent_map.values())

    # Second pass: extract segments from content sheets
    seg_idx = 0
    for sheet_name in wb.sheetnames:
        is_content = any(kw in sheet_name for kw in CONTENT_SHEET_KEYWORDS)
        is_reference = any(kw in sheet_name for kw in REFERENCE_SHEET_KEYWORDS)

        if is_reference and not is_content:
            continue  # Skip pure reference sheets

        ws = wb[sheet_name]
        if ws.max_row < 2:
            continue

        sheet_rows = []
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True):
            sheet_rows.append([str(c) if c is not None else "" for c in row])

        if not sheet_rows:
            continue

        header = sheet_rows[0]
        ncols = len(header)

        # Find respondent columns (columns with player IDs)
        resp_cols = {}
        for ci, h in enumerate(header):
            if not h:
                continue
            # Check rows for player IDs
            for r in sheet_rows[1:min(6, len(sheet_rows))]:
                if ci < len(r) and r[ci]:
                    val = r[ci].strip()
                    for rid in respondent_map:
                        if rid in val or val in rid:
                            resp_cols[ci] = rid
                            break

        if not resp_cols and ncols >= 3:
            # Try to find columns by checking for player patterns
            for ci in range(min(6, ncols)):
                for r in sheet_rows[1:min(6, len(sheet_rows))]:
                    if ci < len(r) and r[ci]:
                        val = r[ci].strip()
                        if re.match(r'^[PG]\d+', val) or re.match(r'^G\d+', val):
                            resp_cols[ci] = val
                            break

        # Extract segments
        for ci, rid in resp_cols.items():
            if rid not in respondent_map:
                continue
            pid = respondent_map[rid]["speaker_id"]

            for ri, r in enumerate(sheet_rows[1:], 1):
                if ci >= len(r) or not r[ci].strip():
                    continue
                content = r[ci].strip()
                if len(content) < 3:
                    continue

                # Try to find question from row context
                question = ""
                if len(r) > 0 and r[0].strip() and ci != 0:
                    question = r[0].strip()
                    if len(question) > 80:
                        question = extract_core_question(question)

                seg_idx += 1
                all_segments.append({
                    "speaker_id": pid,
                    "speaker_role": "interviewee",
                    "preceding_question": question or f"{sheet_name}",
                    "original_text": content,
                    "source_file": f"{source_file}/{sheet_name}",
                })

    return all_respondents, all_segments


# ============================================================
# FILE DISPATCHER
# ============================================================

def _auto_detect_moderator(doc, speaker_patterns):
    """Scan first 40 paragraphs to find the moderator by content heuristics."""
    paragraphs = [p.text.strip() for p in doc.paragraphs[:40]]
    labeled = []
    i = 0
    while i < len(paragraphs):
        text = paragraphs[i]
        if not text:
            i += 1
            continue
        for pat in speaker_patterns:
            m = pat.match(text)
            if m:
                speaker = m.group(1).strip()
                content = m.group(2).strip() if m.lastindex and m.lastindex >= 2 else ""
                if not content and i + 1 < len(paragraphs):
                    next_text = paragraphs[i + 1]
                    is_new = any(p2.match(next_text) for p2 in speaker_patterns)
                    if not is_new:
                        content = next_text
                        i += 1
                if content:
                    labeled.append((speaker, content))
                break
        i += 1

    for speaker, content in labeled:
        lower = content.lower()
        if any(phrase in lower for phrase in MOD_PHRASES):
            return speaker
    return None

def detect_format(filepath, rel_path):
    """Detect file format and return (format_type, extra_params)."""
    fname = os.path.basename(filepath)
    ext = os.path.splitext(fname)[1].lower()

    if ext == '.xlsx':
        return ('xlsx', {})

    doc = Document(filepath)

    # Check if it's paragraph-only (no tables)
    if not doc.tables:
        if "美国HD端射击市场用户细分研究" in rel_path:
            # Auto-detect moderator instead of hardcoding SPEAKER_05
            auto_mod = _auto_detect_moderator(doc, SPEAKER_PATTERNS)
            if auto_mod:
                return ('paragraph', {"moderator_names": {auto_mod}})
            # Fallback: try common moderator patterns
            return ('paragraph', {"moderator_names": {"SPEAKER_05", "SPEAKER_04", "SPEAKER_03", "SPEAKER_06"}})
        if "美国HD端用户生态与决策链路研究" in rel_path:
            return ('paragraph', {"moderator_names": {"Moderator", "moderator"}})
        if "漫威争锋" in rel_path and "文字转录" in rel_path:
            # Text transcription: casual conversation, not structured Q&A
            return ('paragraph', {"moderator_names": {"STF#说话人01"}, "is_text_transcript": True})
        if "绝地潜兵2" in rel_path:
            first_para = doc.paragraphs[0].text.strip() if doc.paragraphs else ""
            if first_para == "WEBVTT":
                return ('webvtt', {})
            return ('paragraph', {"moderator_names": {"主持人"}})
        if "搜打撤品类研究/海外" in rel_path:
            if "座谈会" in rel_path:
                return ('paragraph', {"moderator_names": {"Moderator", "moderator"}})
            else:
                return ('paragraph', {"moderator_names": {"M", "男"}})
        return ('paragraph', {"moderator_names": None})

    # Has tables
    num_tables = len(doc.tables)

    if num_tables >= 2:
        t0 = doc.tables[0]
        if len(t0.columns) >= 5:
            return ('focus_group_2table', {})
        return ('focus_group_1table', {})

    if num_tables == 1:
        table = doc.tables[0]
        if len(table.columns) != 2:
            return ('unknown', {})

        speaker_prefixes = set()
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                user_text = cells[1]
                m = re.match(r'^(G\d+-[A-Za-z0-9_]+)[：:]', user_text)
                if m:
                    speaker_prefixes.add(m.group(1))
                    continue
                m = re.match(r'^(\d+-[^：:]+)[：:]', user_text)
                if m:
                    speaker_prefixes.add(m.group(1))
                    continue
                m = re.match(r'^(\d+)[：:]', user_text)
                if m:
                    speaker_prefixes.add(m.group(1))
                    continue
                m = re.match(r'^([一-鿿]{2,4})[：:]', user_text)
                if m:
                    speaker_prefixes.add(m.group(1))
                    continue
            if len(speaker_prefixes) >= 2:
                break

        if len(speaker_prefixes) >= 2:
            return ('focus_group_1table', {})
        else:
            return ('single_user', {})

    return ('unknown', {})


def process_file(filepath, rel_path):
    """Process one file."""
    ftype, params = detect_format(filepath, rel_path)

    if ftype == 'focus_group_2table':
        return parse_table_focus_group(filepath, rel_path)
    elif ftype == 'focus_group_1table':
        return parse_table_focus_group_1table(filepath, rel_path)
    elif ftype == 'single_user':
        return parse_table_single_user(filepath, rel_path)
    elif ftype == 'player_ability':
        return parse_table_player_ability(filepath, rel_path)
    elif ftype == 'paragraph':
        return parse_paragraph_file(filepath, rel_path, **params)
    elif ftype == 'xlsx':
        return parse_xlsx(filepath, rel_path)
    elif ftype == 'webvtt':
        return parse_webvtt(filepath, rel_path)
    else:
        print(f"  WARNING: Unknown format: {rel_path}")
        return [], []


# ============================================================
# MAIN
# ============================================================

def main():
    all_files = []
    for root, dirs, fnames in os.walk(BASE):
        for f in sorted(fnames):
            if f.endswith(('.docx', '.xlsx')) and not f.startswith('~'):
                full = os.path.join(root, f)
                rel = os.path.relpath(full, BASE)
                all_files.append((full, rel))

    total = len(all_files)
    success = 0
    errors = []
    stats = {"total_segments": 0, "total_respondents": 0, "empty_pq": 0, "short_merged": 0}

    for i, (full, rel) in enumerate(all_files):
        fname = os.path.basename(full)
        print(f"[{i+1}/{total}] {rel} ...", end=" ", flush=True)

        try:
            respondents, segments = process_file(full, rel)

            # Apply post-processing
            respondents, segments = post_process(respondents, segments, rel)

            # Deduplicate respondents by speaker_id
            seen_ids = set()
            unique_respondents = []
            for r in respondents:
                if r["speaker_id"] not in seen_ids:
                    seen_ids.add(r["speaker_id"])
                    unique_respondents.append(r)

            # Count stats
            stats["total_segments"] += len(segments)
            stats["total_respondents"] += len(unique_respondents)
            for s in segments:
                if not s["preceding_question"]:
                    stats["empty_pq"] += 1

            # Build output
            output = {
                "meta": {
                    "version": "v2.2",
                    "processing_date": "2026-08-26",
                    "source_file": rel,
                    "source_type": "座谈会（Focus Group）" if len(unique_respondents) > 1 else "一对一访谈",
                    "participant_count": len(unique_respondents),
                    "segment_count": len(segments),
                    "processing_notes": [
                        "§6.7 默认规则：preceding_question = 直接触发该回答的问题",
                        "例外规则（极少）：真正追问时加「前置问题 → 当前问题」",
                        "点名 → 回到主题提问",
                        "v2.2: PQ 长度上限 200 字符",
                        "v2.2: 空 PQ 回填仅使用真实提问（含?）",
                        "v2.2: 英文对话分类优化",
                        "v2.2: 主持人自动检测",
                        "v2.2: 文字转录特殊处理",
                        "v2.1: 筛选表 Gaming Background 提取",
                        "v2.1: 无筛选表文件 Profile/GB 从 Segment 文本提取",
                        "v2.1: group_code 从文件路径提取",
                        "v2.1: 超短 Segment 合并（≤5字符）",
                        "v2.1: xlsx §6.4 内容/参考 Sheet 区分",
                        "v2.1: display_name 修复",
                    ]
                },
                "respondents": unique_respondents,
                "segments": [
                    {
                        "segment_id": j + 1,
                        "source_file": s.get("source_file", rel),
                        "segment_index": j + 1,
                        "speaker_id": s["speaker_id"],
                        "speaker_role": s.get("speaker_role", "interviewee"),
                        "preceding_question": s["preceding_question"],
                        "original_text": s["original_text"],
                        "cleaned_text": None,
                        "char_count": len(s["original_text"])
                    }
                    for j, s in enumerate(segments)
                ]
            }

            out_rel = os.path.splitext(rel)[0] + ".json"
            out_path = os.path.join(OUT, out_rel)
            os.makedirs(os.path.dirname(out_path), exist_ok=True)
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(output, f, ensure_ascii=False, indent=2)

            print(f"OK ({len(unique_respondents)}R, {len(segments)}S)")
            success += 1

        except Exception as e:
            import traceback
            print(f"ERROR: {e}")
            traceback.print_exc()
            errors.append((rel, str(e)))

    print(f"\n{'='*60}")
    print(f"Done: {success}/{total} files processed successfully")
    print(f"Stats: {stats['total_respondents']} respondents, {stats['total_segments']} segments")
    print(f"Empty PQs after fix: {stats['empty_pq']}")
    if errors:
        print(f"Errors ({len(errors)}):")
        for rel, err in errors:
            print(f"  {rel}: {err}")


if __name__ == "__main__":
    main()