#!/usr/bin/env python3
"""Extract ONE sample segment from 漫威争锋 G1 for review."""
import docx, re, os, json

base = r'C:\Users\鸢尾\Desktop\腾讯-模拟用户\资料\虚拟用户-笔录 for 元培\虚拟用户-笔录 for 元培\漫威争锋中美用户洞察研究'

doc = docx.Document(os.path.join(base, '座谈会笔录-G1.docx'))

# --- Build speaker map from header paragraphs ---
speaker_map = {}
for p in doc.paragraphs:
    text = p.text.strip()
    m = re.match(r'(\d+)\s*(\S.+)', text)
    if m and len(m.group(2)) < 20 and not m.group(2).startswith('实录'):
        speaker_map[m.group(1)] = m.group(2).strip()

# --- Extract segments from table ---
table = doc.tables[0]
segments = []
last_mod_text = None
seg_idx = 0

# Patterns for uninformative content to filter
NOISE_PATTERNS = [
    r'^[对是嗯好行可可以]+[，,。.]?$',
    r'^[没不][有会是知道清楚懂行能]+[，,。.]?$',
    r'^(Cool|Yeah|Yep|Nope|Yes|No|Okay|Ok|Sure|Right|Gotcha|Perfect|Fantastic|Great|Nice|Thanks|Thank\s*you|Sorry|Fine|Alright|Absolutely|Exactly|Definitely|Totally|Indeed|Correct|Fair|True|Maybe|Perhaps|Probably)$',
    r'^[Mm]+[hm]+[，,。.]?$',
    r'^[啊哦嗯呃唉哎哟嘿]$',
    r'^[?!？！.。，,]+$',
]

def is_noise(text):
    """Check if text is uninformative filler."""
    t = text.strip()
    if len(t) <= 2:
        return True
    for pat in NOISE_PATTERNS:
        if re.match(pat, t, re.IGNORECASE):
            return True
    # Very short text without substantive content
    if len(t) <= 8 and not re.search(r'[《》\w]{2,}|【.*?】', t):
        return True
    return False

def is_mod_name_only(text):
    """Check if moderator text is just a name callout."""
    t = text.strip()
    t = re.sub(r'^[Mm]\s*[：:]\s*', '', t)
    # Just a name or "name+说一下/介绍一下/你来说说"
    if re.match(r'^[一-鿿A-Za-z0-9_\-]{1,10}[。，,.]?\s*$', t):
        return True
    if re.match(r'^[一-鿿A-Za-z0-9_\-]{1,10}[说一下介绍一下补充一下来说说].*$', t):
        return True
    if len(t) <= 6:
        return True
    return False

for row in table.rows:
    mod_text = row.cells[0].text.strip()
    int_text = row.cells[1].text.strip()

    # Update moderator context
    if mod_text:
        mod_clean = re.sub(r'^[Mm]\s*[：:]\s*', '', mod_text).strip()
        if mod_clean and not is_mod_name_only(mod_text):
            last_mod_text = mod_clean

    if not int_text:
        continue

    # Parse interviewee
    m = re.match(r'(\d+)\s*[：:]\s*(.+)', int_text, re.DOTALL)
    if not m:
        continue

    num = m.group(1)
    text = m.group(2).strip()
    speaker_id = speaker_map.get(num, f'P{num}')

    # Filter noise
    if is_noise(text):
        continue

    segments.append({
        'source_file': '漫威争锋中美用户洞察研究/座谈会笔录-G1.docx',
        'segment_index': seg_idx,
        'speaker_id': speaker_id,
        'speaker_role': 'interviewee',
        'preceding_question': last_mod_text,
        'original_text': text,
        'cleaned_text': None,
        'char_count': None,
        'annotation': {}
    })
    seg_idx += 1

# Write first 3 segments to file for review
outpath = os.path.join(r'data\sheets_data', '_sample_output.json')
with open(outpath, 'w', encoding='utf-8') as f:
    json.dump(segments[:3], f, ensure_ascii=False, indent=2)
print(f'Written to {outpath}')
print(f'Total: {len(segments)} segments (filtered from 613)')