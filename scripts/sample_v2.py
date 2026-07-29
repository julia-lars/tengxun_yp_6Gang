#!/usr/bin/env python3
"""Extract ONE sample: 1 respondent + 1 segment from 漫威争锋 G1."""
import docx, re, os, json

base = r'C:\Users\鸢尾\Desktop\腾讯-模拟用户\资料\虚拟用户-笔录 for 元培\虚拟用户-笔录 for 元培\漫威争锋中美用户洞察研究'

doc = docx.Document(os.path.join(base, '座谈会笔录-G1.docx'))

# --- Build speaker map from header ---
speaker_map = {}
for p in doc.paragraphs:
    text = p.text.strip()
    m = re.match(r'(\d+)\s*(\S.+)', text)
    if m and len(m.group(2)) < 20 and not m.group(2).startswith('实录'):
        speaker_map[m.group(1)] = m.group(2).strip()

# --- Extract all rows from table ---
rows = []
table = doc.tables[0]
for row in table.rows:
    mod = row.cells[0].text.strip()
    intv = row.cells[1].text.strip()
    if intv:
        m = re.match(r'(\d+)\s*[：:]\s*(.+)', intv, re.DOTALL)
        if m:
            num = m.group(1)
            text = m.group(2).strip()
            sid = speaker_map.get(num, f'P{num}')
            mod_clean = re.sub(r'^[Mm]\s*[：:]\s*', '', mod).strip() if mod else ''
            rows.append({'speaker_id': sid, 'mod': mod_clean, 'text': text, 'num': num})

# ============================================================
# Separate: first 8 rows = self-intro (respondents), rest = game discussion (segments)
# ============================================================

# --- RESPONDENTS: self-intro rows ---
respondents = {}
for r in rows[:15]:  # first ~15 rows contain intros
    sid = r['speaker_id']
    if sid not in respondents:
        respondents[sid] = {
            'source_file': '漫威争锋中美用户洞察研究/座谈会笔录-G1.docx',
            'speaker_id': sid,
            'display_name': None,
            'group_code': 'G1',
            'background': {
                'profile': None,
                'game_experience': None,
                'game_experience_summary': None
            },
            '_intro_texts': []
        }
    respondents[sid]['_intro_texts'].append(r['text'])

# Merge intro texts into profile
for sid, resp in respondents.items():
    full = '；'.join(resp['_intro_texts'])
    resp['background']['profile'] = full
    resp['background']['game_experience'] = full
    resp['background']['game_experience_summary'] = full
    del resp['_intro_texts']

# --- SEGMENTS: game discussion rows (skip first ~8 intro rows) ---
segments = []
last_mod = None
seg_idx = 0

NOISE_PATTERNS = [
    r'^[对是嗯好行可可以]+[，,。.]?$',
    r'^[没不][有会是知道清楚懂行能]+[，,。.]?$',
    r'^(Cool|Yeah|Yep|Nope|Yes|No|Okay|Ok|Sure|Right|Gotcha|Perfect|Fantastic|Great|Nice|Thanks|Sorry|Fine|Alright|Absolutely|Exactly|Definitely|Totally|Indeed|Correct|Fair|True|Maybe|Perhaps|Probably)$',
    r'^[Mm]+[hm]+[，,。.]?$',
    r'^[啊哦嗯呃唉哎哟嘿]$',
    r'^[?!？！.。，,]+$',
]

def is_noise(text):
    t = text.strip()
    if len(t) <= 2:
        return True
    for pat in NOISE_PATTERNS:
        if re.match(pat, t, re.IGNORECASE):
            return True
    if len(t) <= 8 and not re.search(r'[《》\w]{2,}|【.*?】', t):
        return True
    return False

def is_name_only(text):
    t = re.sub(r'^[Mm]\s*[：:]\s*', '', text).strip()
    if re.match(r'^[一-鿿A-Za-z0-9_\-]{1,10}[。，,.]?\s*$', t):
        return True
    if re.match(r'^[一-鿿A-Za-z0-9_\-]{1,10}[说一下介绍一下补充一下来说说].*$', t):
        return True
    return len(t) <= 6

for r in rows[8:]:  # skip intro rows
    if r['mod'] and not is_name_only(r['mod']):
        last_mod = r['mod']

    if is_noise(r['text']):
        continue

    segments.append({
        'source_file': '漫威争锋中美用户洞察研究/座谈会笔录-G1.docx',
        'segment_index': seg_idx,
        'speaker_id': r['speaker_id'],
        'speaker_role': 'interviewee',
        'preceding_question': last_mod,
        'original_text': r['text'],
        'cleaned_text': None,
        'char_count': None,
        'annotation': {}
    })
    seg_idx += 1

# --- Output to file ---
out = {
    'respondent_sample': list(respondents.values())[0],
    'segment_sample': segments[0],
    'total_respondents': len(respondents),
    'total_segments': len(segments)
}
outpath = os.path.join(r'data\sheets_data', '_sample_v2.json')
with open(outpath, 'w', encoding='utf-8') as f:
    json.dump(out, f, ensure_ascii=False, indent=2)
print(f'Written to {outpath}')
print(f'Total respondents: {len(respondents)}, Total segments: {len(segments)}')