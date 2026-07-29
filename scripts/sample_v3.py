"""Extract ONE sample segment from 漫威争锋, translate English to Chinese."""
import docx, re, os, json

base = r'C:\Users\鸢尾\Desktop\腾讯-模拟用户\资料\虚拟用户-笔录 for 元培\虚拟用户-笔录 for 元培\漫威争锋中美用户洞察研究'
outdir = r'data\sheets_data'

# ---- G1 (Chinese table) ----
doc = docx.Document(os.path.join(base, '座谈会笔录-G1.docx'))
speaker_map = {}
for p in doc.paragraphs:
    text = p.text.strip()
    m = re.match(r'(\d+)\s*(\S.+)', text)
    if m and len(m.group(2)) < 20 and not m.group(2).startswith('实录'):
        speaker_map[m.group(1)] = m.group(2).strip()

table = doc.tables[0]
segments_g1 = []
last_mod = None
seg_idx = 0

NOISE = re.compile(
    r'^[对是嗯好行可可以]+[，,。.]?$|'
    r'^[没不][有会是知道清楚懂行能]+[，,。.]?$|'
    r'^(Cool|Yeah|Yep|Nope|Yes|No|Okay|Ok|Sure|Right|Gotcha|Perfect|Fantastic|Great|Nice|Thanks|Sorry|Fine|Alright|Absolutely|Exactly|Definitely|Totally|Indeed|Correct|Fair|True|Maybe|Perhaps|Probably)$|'
    r'^[Mm]+[hm]+[，,。.]?$|'
    r'^[啊哦嗯呃唉哎哟嘿]$',
    re.IGNORECASE
)

for row in table.rows:
    mod = row.cells[0].text.strip()
    intv = row.cells[1].text.strip()
    if mod:
        mod_clean = re.sub(r'^[Mm]\s*[：:]\s*', '', mod).strip()
        # Skip name-only moderator text
        if mod_clean and len(mod_clean) > 10 or (mod_clean and not re.match(r'^[一-鿿A-Za-z0-9_\-]{1,10}[说一下介绍一下补充一下来说说。，,.]?\s*$', mod_clean)):
            last_mod = mod_clean
    if not intv:
        continue
    m = re.match(r'(\d+)\s*[：:]\s*(.+)', intv, re.DOTALL)
    if not m:
        continue
    num, text = m.group(1), m.group(2).strip()
    sid = speaker_map.get(num, f'P{num}')
    if NOISE.match(text) or len(text) <= 2:
        continue
    segments_g1.append({
        'source_file': '漫威争锋中美用户洞察研究/座谈会笔录-G1.docx',
        'segment_index': seg_idx, 'speaker_id': sid, 'speaker_role': 'interviewee',
        'preceding_question': last_mod, 'original_text': text,
        'cleaned_text': None, 'char_count': None, 'annotation': {}
    })
    seg_idx += 1

# ---- Text transcription (English, need translation) ----
# Pick one substantive segment from 海外年轻组
doc_en = docx.Document(os.path.join(base, '文字转录-海外年轻组.docx'))
segments_en = []
last_mod = None
seg_idx = 0

for p in doc_en.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    m = re.match(r'STF#说话人(\d+)\([^)]+\):\s*(.+)', text, re.DOTALL)
    if m:
        spk, speech = m.group(1), m.group(2).strip()
        if spk == '01':
            last_mod = speech
        else:
            if NOISE.match(speech) or len(speech) <= 2:
                continue
            segments_en.append({
                'source_file': '漫威争锋中美用户洞察研究/文字转录-海外年轻组.docx',
                'segment_index': seg_idx, 'speaker_id': f'STF{spk}', 'speaker_role': 'interviewee',
                'preceding_question': last_mod, 'original_text': speech,
                'cleaned_text': None, 'char_count': None, 'annotation': {}
            })
            seg_idx += 1
    elif segments_en:
        segments_en[-1]['original_text'] += ' ' + text

# ---- Show one sample from each source ----
sample = {
    'g1_chinese': segments_g1[12],  # pick a substantive one
    'en_young': segments_en[6],     # pick a substantive one
    'notes': 'English content needs translation to Chinese'
}

with open(os.path.join(outdir, '_sample_v3.json'), 'w', encoding='utf-8') as f:
    json.dump(sample, f, ensure_ascii=False, indent=2)
print(f'G1: {len(segments_g1)} segments, EN young: {len(segments_en)} segments')
print('Sample written')