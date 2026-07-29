"""Show ONE respondent + ONE segment from 漫威争锋 G1."""
import docx, re, os, json

base = r'C:\Users\鸢尾\Desktop\腾讯-模拟用户\资料\虚拟用户-笔录 for 元培\虚拟用户-笔录 for 元培\漫威争锋中美用户洞察研究'

doc = docx.Document(os.path.join(base, '座谈会笔录-G1.docx'))

# --- Speaker map from header ---
speaker_map = {}
for p in doc.paragraphs:
    text = p.text.strip()
    m = re.match(r'(\d+)\s*(\S.+)', text)
    if m and len(m.group(2)) < 20 and not m.group(2).startswith('实录'):
        speaker_map[m.group(1)] = m.group(2).strip()

# --- Parse all rows ---
all_rows = []
table = doc.tables[0]
for row in table.rows:
    mod = row.cells[0].text.strip()
    intv = row.cells[1].text.strip()
    if intv:
        m = re.match(r'(\d+)\s*[：:]\s*(.+)', intv, re.DOTALL)
        if m:
            num, text = m.group(1), m.group(2).strip()
            sid = speaker_map.get(num, f'P{num}')
            mod_clean = re.sub(r'^[Mm]\s*[：:]\s*', '', mod).strip() if mod else ''
            all_rows.append({'speaker_id': sid, 'mod': mod_clean, 'text': text, 'num': num})

# ------ RESPONDENTS: self-intro rows (first ~8 speakers' first lines) ------
# The moderator says "大家先个人自我介绍...先顺时针开始，从G1-LJH开始"
# Rows 0-8 are self-introductions
intro_end = 9  # Rows 0-8 are self-introductions for all 8 speakers
respondents = {}
for r in all_rows[:intro_end]:
    sid = r['speaker_id']
    if sid not in respondents:
        respondents[sid] = {
            'source_file': '漫威争锋中美用户洞察研究/座谈会笔录-G1.docx',
            'speaker_id': sid,
            'display_name': None,
            'group_code': 'G1',
            'background': {
                'profile': None,
                'game_experience': None,
                'game_experience_summary': None
            }
        }
    respondents[sid]['background']['profile'] = r['text']

# ------ SEGMENTS: game discussion (after intro) ------
NOISE = re.compile(
    r'^[对是嗯好行可可以]+[，,。.]?$|'
    r'^[没不][有会是知道清楚懂行能]+[，,。.]?$|'
    r'^(Cool|Yeah|Yep|Nope|Yes|No|Okay|Ok|Sure|Right|Gotcha|Perfect|Fantastic|Great|Nice|Thanks|Sorry|Fine|Alright|Absolutely|Exactly|Definitely|Totally|Indeed|Correct|Fair|True|Maybe|Perhaps|Probably)$|'
    r'^[Mm]+[hm]+[，,。.]?$|'
    r'^[啊哦嗯呃唉哎哟嘿]$',
    re.IGNORECASE
)

segments = []
last_mod = None
seg_idx = 0

for r in all_rows[intro_end:]:
    if r['mod'] and len(r['mod']) > 10:
        last_mod = r['mod']
    if NOISE.match(r['text']) or len(r['text']) <= 2:
        continue
    segments.append({
        'source_file': '漫威争锋中美用户洞察研究/座谈会笔录-G1.docx',
        'segment_index': seg_idx, 'speaker_id': r['speaker_id'], 'speaker_role': 'interviewee',
        'preceding_question': last_mod, 'original_text': r['text'],
        'cleaned_text': None, 'char_count': None, 'annotation': {}
    })
    seg_idx += 1

# --- Output ---
sample = {
    'respondent_sample': list(respondents.values())[0],
    'segment_sample': segments[0],
    'stats': f'{len(respondents)} respondents, {len(segments)} segments'
}
with open(r'data\sheets_data\_sample_v4.json', 'w', encoding='utf-8') as f:
    json.dump(sample, f, ensure_ascii=False, indent=2)
print('Done')